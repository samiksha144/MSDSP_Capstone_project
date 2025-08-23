import io
import os
import re
import textwrap
from datetime import datetime
from typing import List, Tuple, Dict, Optional
from docx.shared import Pt, Inches  # add Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

import pandas as pd
import streamlit as st



# === DB audits integration (writes; no admin table on this page) ===
from db_repo import (
    ensure_audit_schema,
    add_audit as db_add_audit,
    DBError,
)

# ---------------------------
# Page config
# ---------------------------
st.set_page_config(page_title="REGDOCGPT", page_icon="📄", layout="wide")

# ---------------------------
# Page background
# ---------------------------
def set_grey_bg(color: str = "#000000"):
    st.markdown(
        f"""
        <style>
          [data-testid="stAppViewContainer"] {{
            background: {color} !important;
          }}
          [data-testid="stHeader"] {{
            background: {color} !important;
            z-index: 2 !important;
            border-bottom: 1px solid #e5e7eb;
          }}
          .block-container {{
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-radius: 16px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.05);
            padding-top: 1.2rem;
            padding-bottom: 1.2rem;
          }}
        </style>
        """,
        unsafe_allow_html=True,
    )

set_grey_bg("#000000")

# Hide Streamlit's top-right Deploy/status toolbar
def _hide_deploy_button():
    st.markdown(
        """
        <style>
        div[data-testid="stToolbar"] { display: none !important; }
        div[data-testid="stStatusWidget"] { display: none !important; }
        </style>
        """,
        unsafe_allow_html=True,
    )

_hide_deploy_button()

# ======== Require login & access session account ========
acct = st.session_state.get("account")
if not acct:
    st.switch_page("pages/login.py")

# --- Profile menu actions via query params ---
action = st.query_params.get("action")
if isinstance(action, list):
    action = action[0] if action else None

def _clear_action_param():
    try:
        del st.query_params["action"]
    except Exception:
        pass

if action == "logout":
    st.session_state.pop("account", None)
    _clear_action_param()
    st.switch_page("pages/login.py")

elif action == "my_docs":
    st.session_state["jump_to_uploads"] = True
    _clear_action_param()
    st.rerun()

def show_profile_menu():
    """Fixed top-right profile dropdown (no code block rendering)."""
    name = acct.get("username", "user")
    uid  = acct.get("id", "—")

    html = f"""
<style>
[data-testid="stHeader"] {{ background: transparent; z-index: 0 !important; }}
.profile-menu {{ position: fixed; top: 10px; right: 16px; z-index: 999999 !important;
  font-family: ui-sans-serif,-apple-system,Segoe UI,Roboto,Helvetica,Arial; }}
.profile-menu details {{ position: relative; }}
.profile-menu summary {{ list-style: none; outline: none; cursor: pointer; display: flex; align-items: center; gap: 10px;
  padding: 6px 12px; background: rgba(255,255,255,0.98); border: 1px solid #e5e7eb; border-radius: 9999px;
  box-shadow: 0 8px 22px rgba(0,0,0,.10); user-select: none; }}
.profile-menu summary::-webkit-details-marker {{ display:none; }}
.profile-menu .name {{ font-weight: 700; color:#111827; line-height: 1; }}
.profile-menu .meta {{ font-size: 12px; color:#6b7280; margin-top:-2px; }}
.profile-menu .menu {{ position: absolute; right: 0; top: 100%; margin-top: 8px; min-width: 240px; background: #fff;
  border: 1px solid #e5e7eb; border-radius: 12px; box-shadow: 0 16px 40px rgba(0,0,0,.18); padding: 6px; }}
.profile-menu a {{ display: block; padding: 10px 12px; border-radius: 8px; text-decoration: none; color: #111827; }}
.profile-menu a:hover {{ background: #f3f4f6; }}
</style>
<div class="profile-menu">
<details>
<summary title="Open profile menu">
  <span style="font-size:18px;">👤</span>
  <div>
    <div class="name">{name}</div>
    <div class="meta">ID: {uid}</div>
  </div>
  <span style="margin-left:6px;">▾</span>
</summary>
<div class="menu">
  <a href="?action=logout" target="_self">Sign out</a>
</div>
</details>
</div>
"""
    st.markdown(html, unsafe_allow_html=True)

show_profile_menu()

# ---------------------------
# Optional deps
# ---------------------------
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None
try:
    import docx2txt
except Exception:
    docx2txt = None

# PDF & DOCX builders
try:
    from reportlab.lib.pagesizes import LETTER  # type: ignore
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    from reportlab.lib import colors  # type: ignore
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    Pt = None
    WD_ALIGN_PARAGRAPH = None
    DOCX_AVAILABLE = False

# ---- OpenAI (new SDK preferred; legacy supported) ----
_OPENAI_AVAILABLE = False
try:
    from openai import OpenAI  # >=1.0
    _OPENAI_AVAILABLE = True
except Exception:
    try:
        import openai  # legacy 0.x
        _OPENAI_AVAILABLE = True
    except Exception:
        _OPENAI_AVAILABLE = False

# ---------------------------
# DB bootstrap (audits table + procs) — run once
# ---------------------------
@st.cache_resource(show_spinner=False)
def _bootstrap_audit_schema() -> Optional[str]:
    try:
        ensure_audit_schema()
        return None
    except Exception as e:
        return f"{e}"

_bootstrap_err = _bootstrap_audit_schema()
if _bootstrap_err:
    st.warning(f"Audit DB bootstrap warning: {_bootstrap_err}")

# ---------------------------
# Utilities
# ---------------------------
def now_iso() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def _safe_int(x) -> Optional[int]:
    try:
        return int(x)
    except Exception:
        return None

# ---- Unified add_audit wrapper ----
def add_audit(actor: str, event: str, detail: str = "") -> None:
    acct_obj = st.session_state.get("account") or {}
    role = (acct_obj.get("role") or "user").lower()
    uid = _safe_int(acct_obj.get("id"))

    admin_id = uid if role == "admin" else None
    user_id = uid if role != "admin" else None

    try:
        db_add_audit(
            actor=actor,
            actor_role=role if role in ("user", "admin", "system", "assistant") else None,
            admin_id=admin_id,
            user_id=user_id,
            event=event,
            detail=detail,
        )
        return
    except TypeError:
        try:
            db_add_audit(
                actor=actor,
                user_id=user_id if user_id is not None else admin_id,
                event=event,
                detail=detail,
            )
            return
        except Exception as e:
            _append_audit_memory(actor, uid, event, f"(DB v1 fail) {detail} | {e}")
    except Exception as e:
        _append_audit_memory(actor, uid, event, f"(DB v2 fail) {detail} | {e}")

def _append_audit_memory(actor: str, uid: Optional[int], event: str, detail: str) -> None:
    if "audit" not in st.session_state:
        st.session_state.audit = []
    st.session_state.audit.insert(0, {
        "Timestamp": now_iso(),
        "Actor": actor,
        "UserID": str(uid) if uid is not None else "—",
        "Event": event,
        "Detail": detail,
    })

def summarize_heuristic(text: str) -> str:
    cleaned = " ".join((text or "").split())
    if not cleaned:
        return "The SOP document outlines the procedures for regulatory compliance and quality management."
    return (cleaned[:1200] + "…") if len(cleaned) > 1200 else cleaned

# ---------------------------
# PDF helpers
# ---------------------------
def _pdf_story_heading(text: str, styles):
    return Paragraph(f"<b>{text}</b>", styles["Heading4"])

def _pdf_story_paras(text: str, styles):
    parts = []
    for para in (text or "").split("\n"):
        parts.append(Paragraph(para.strip() or "&nbsp;", styles["BodyText"]))
    return parts

def make_pdf_from_template(sop_fields: Dict[str, str], meta: Dict[str, str]) -> bytes:
    """Render PDF with the header table layout matching the 1st screenshot."""
    bio = io.BytesIO()
    if not REPORTLAB_AVAILABLE:
        # Fallback minimal text
        bio.write((meta.get("title", "Standard Operating Procedure") + "\n\n").encode("utf-8"))
        for k, v in sop_fields.items():
            bio.write((f"{k}\n{v}\n\n").encode("utf-8"))
        return bio.getvalue()

    doc = SimpleDocTemplate(bio, pagesize=LETTER, topMargin=36, bottomMargin=36, leftMargin=48, rightMargin=48)
    styles = getSampleStyleSheet()
    story = []

    # Header title
    story.append(Paragraph("<b>STANDARD OPERATING PROCEDURE (SOP)</b>", styles["Title"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph(meta.get("organization", ""), styles["Heading3"]))
    story.append(Spacer(1, 12))

    # Header table like image #1
    data = [
        ["ASGS Pharmaceuticals", "Department:", meta.get("department", ""), f"SOP No.: {meta.get('sop_no','')}"],
        ["", "Area:", meta.get("area", ""), f"Effective Date: {meta.get('effective_date','')}"],
        ["", "Title:", meta.get("title",""), f"Review Date: {meta.get('review_date','')}"],
        ["", "Prepared By", "Reviewed By", "Approved By"],
        ["", "Signature", "Signature", "Signature"],
        ["", "Name", "Name", "Name"],
        ["", "Date", "Date", "Date"],
    ]
    tbl = Table(data, colWidths=[120, 150, 170, 150])
    ts = TableStyle([
        ("GRID", (0,0), (-1,-1), 0.8, colors.black),
        ("SPAN", (0,0), (0,6)),           # left merged cell
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN", (0,0), (0,6), "CENTER"),
        ("FONTNAME", (0,0), (0,6), "Helvetica-Bold"),
        ("FONTNAME", (1,0), (3,0), "Helvetica-Bold"),
        ("FONTNAME", (1,1), (3,1), "Helvetica-Bold"),
        ("FONTNAME", (1,2), (3,2), "Helvetica-Bold"),
        ("FONTNAME", (1,3), (3,3), "Helvetica-Bold"),
    ])
    tbl.setStyle(ts)
    story.append(tbl)
    story.append(Spacer(1, 18))

    # Sections
    section_order = [
        ("1. Purpose", sop_fields.get("1. Purpose")),
        ("2. Scope", sop_fields.get("2. Scope")),
        ("3. Responsibilities", sop_fields.get("3. Responsibilities")),
        ("4. Definitions", sop_fields.get("4. Definitions")),
        ("5. References", sop_fields.get("5. References")),
        ("6. Procedure", sop_fields.get("6. Procedure")),
        ("6.1 Materials & Equipment", sop_fields.get("6.1 Materials & Equipment")),
        ("6.2 Stepwise Procedure", sop_fields.get("6.2 Stepwise Procedure")),
        ("7. Safety Precautions", sop_fields.get("7. Safety Precautions")),
        ("8. Training Requirements", sop_fields.get("8. Training Requirements")),
        ("9. Change Control", sop_fields.get("9. Change Control")),
        ("10. Records & Documentation", sop_fields.get("10. Records & Documentation")),
        ("11. Revision History", sop_fields.get("11. Revision History")),
    ]
    for title, body in section_order:
        story.append(_pdf_story_heading(title, styles))
        story.extend(_pdf_story_paras(body or "", styles))
        story.append(Spacer(1, 10))

    doc.build(story)
    return bio.getvalue()

# ---------------------------
# DOCX helpers
# ---------------------------
def _add_heading(doc, text, size=13, bold=True, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    if Pt:
        r.font.size = Pt(size)
    r.bold = bold

def _add_paragraphs(doc, text: str):
    for para in (text or "").split("\n"):
        doc.add_paragraph(para.strip())

def _add_training_table(doc, rows: int = 10):
    table = doc.add_table(rows=rows + 1, cols=4)
    table.style = "Table Grid"
    hdr = ["S.No.", "Trainee Name", "Trained on", "Sign and Date"]
    for j, h in enumerate(hdr):
        cell = table.cell(0, j)
        cell.text = h

    # set approximate widths (looks good in Word; adjust if needed)
    widths = [Inches(0.8), Inches(2.2), Inches(2.0), Inches(2.0)]
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = w

    # empty body rows
    for i in range(1, rows + 1):
        table.cell(i, 0).text = str(i)

def _add_revision_history_table(doc, initial_version: str = "v1.0",
                                date_str: str = "", initial_desc: str = "Initial issue.",
                                rows: int = 8):
    table = doc.add_table(rows=rows + 1, cols=4)
    table.style = "Table Grid"
    hdr = ["Version", "Date", "Description of Change", "Approved By"]
    for j, h in enumerate(hdr):
        table.cell(0, j).text = h

    widths = [Inches(1.0), Inches(1.3), Inches(3.7), Inches(1.6)]
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = w

    # Prefill first row
    if rows >= 1:
        table.cell(1, 0).text = initial_version
        table.cell(1, 1).text = date_str
        table.cell(1, 2).text = initial_desc
        table.cell(1, 3).text = ""

def _clean_section_body(heading: str, body: str) -> str:
    """Strip markdown and drop a leading line that duplicates the heading."""
    txt = _strip_markdown(body or "")
    lines = [ln.strip() for ln in txt.splitlines()]

    # drop leading empties
    while lines and not lines[0]:
        lines.pop(0)

    # normalize for comparison (ignore spaces, punctuation)
    def _norm(s: str) -> str:
        return re.sub(r'[\s:.\-]+', ' ', s).strip().lower()

    if lines and _norm(lines[0]) == _norm(heading):
        lines.pop(0)

    return ("\n".join(lines)).strip()


def _add_footer_with_form_and_page(doc, form_no: str = "QA 01.04.02/14"):
    """
    Footer: left = Form No., right = PAGE field
    Works on all pages of the .docx.
    """
    section = doc.sections[0]
    footer = section.footer

    # REQUIRED: width when adding a table to header/footer
    avail_width = section.page_width - section.left_margin - section.right_margin
    tbl = footer.add_table(rows=1, cols=2, width=avail_width)

    # (Optional) borderless look — comment these lines if you like "Table Grid"
    tbl.style = None
    for r in tbl.rows:
        for c in r.cells:
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            for tag in ("w:tcBorders",):
                el = tcPr.find(qn(tag))
                if el is not None:
                    tcPr.remove(el)

    # Set approximate column widths (65% / 35%)
    try:
        tbl.columns[0].width = int(avail_width * 0.65)
        tbl.columns[1].width = int(avail_width * 0.35)
    except Exception:
        # Some python-docx builds ignore column width; it's fine — alignment still works
        pass

    # LEFT cell: Form No.
    left_p = tbl.cell(0, 0).paragraphs[0]
    left_p.text = f"Form No.: {form_no}"

    # RIGHT cell: PAGE field (right-aligned)
    right_p = tbl.cell(0, 1).paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")  # Word renders current page number
    right_p._p.append(fld)


def _add_page_border(doc, border_size: int = 24, border_color: str = "000000", space: int = 24):
    """
    Adds a full-page border (all four sides) to every section.
    border_size: thickness in eighths of a point (24 = 3pt, 48 = 6pt).
    border_color: hex RGB string, e.g. "000000".
    space: gap between border and page content, in eighths of a point.
    """
    for section in doc.sections:
        sectPr = section._sectPr

        # Remove an existing <w:pgBorders> if present (avoids duplicates/no-ops)
        existing = sectPr.find(qn('w:pgBorders'))
        if existing is not None:
            sectPr.remove(existing)

        pgBorders = OxmlElement('w:pgBorders')
        # Important: draw relative to the page, and place in front
        pgBorders.set(qn('w:offsetFrom'), 'page')
        pgBorders.set(qn('w:zOrder'), 'front')  # ensures the border is visible

        for side in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), str(border_size))   # thickness (1 pt = 8 units)
            b.set(qn('w:space'), str(space))      # gap to text
            b.set(qn('w:color'), border_color)
            pgBorders.append(b)

        sectPr.append(pgBorders)

# ---------------------------
# Full SOP DOCX builder
# ---------------------------
def make_docx_from_template(sop_fields: Dict[str, str], meta: Dict[str, str]) -> bytes:
    if not DOCX_AVAILABLE:
        buf = io.BytesIO()
        title = sop_fields.get("Title") or meta.get("title") or "Standard Operating Procedure"
        buf.write(f"STANDARD OPERATING PROCEDURE (SOP)\n{meta.get('organization','')}\n\n{title}\n\n".encode("utf-8"))
        for key in SECTION_KEYS:
            if key == "Title":
                continue
            buf.write((f"{key}\n{_strip_markdown(sop_fields.get(key, ''))}\n\n").encode("utf-8"))
        # Simple textual placeholders for the ending tables
        buf.write(("Training Requirements\nS.No. | Trainee Name | Trained on | Sign and Date\n").encode("utf-8"))
        buf.write(("Revision History\nVersion | Date | Description of Change | Approved By\n").encode("utf-8"))
        return buf.getvalue()

    doc = Document()

    # Header
    _add_heading(doc, "STANDARD OPERATING PROCEDURE (SOP)", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_heading(doc, meta.get("organization", ""), size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # 7x4 header table (as in your file)
    tbl = doc.add_table(rows=7, cols=4)
    tbl.style = "Table Grid"
    brand_cell = tbl.cell(0, 0).merge(tbl.cell(6, 0))
    brand_cell.text = "ASGS\nPharmaceuticals"

    tbl.cell(0, 1).text = "Department:"
    tbl.cell(0, 2).text = meta.get("department", "")
    tbl.cell(0, 3).text = f"SOP No.: {meta.get('sop_no','')}"
    tbl.cell(1, 1).text = "Area:"
    tbl.cell(1, 2).text = meta.get("area", "")
    tbl.cell(1, 3).text = f"Effective Date: {meta.get('effective_date','')}"
    tbl.cell(2, 1).text = "Title:"
    # prefer SOP Title from fields; fallback to meta
    sop_title_text = sop_fields.get("Title") or meta.get("title", "")
    tbl.cell(2, 2).text = _strip_markdown(sop_title_text)
    tbl.cell(2, 3).text = f"Review Date: {meta.get('review_date','')}"
    tbl.cell(3, 1).text = "Prepared By"
    tbl.cell(3, 2).text = "Reviewed By"
    tbl.cell(3, 3).text = "Approved By"
    for i, label in enumerate(["Signature", "Name", "Date"], start=4):
        tbl.cell(i, 1).text = label
        tbl.cell(i, 2).text = label
        tbl.cell(i, 3).text = label

    doc.add_paragraph()

    # Body sections with 1.0/2.0 numbering and cleaned text
    section_order = [
        ("Title", sop_title_text),
        ("1.0 Purpose", sop_fields.get("1. Purpose", "")),
        ("2.0 Scope", sop_fields.get("2. Scope", "")),
        ("3.0 Responsibilities", sop_fields.get("3. Responsibilities", "")),
        ("4.0 Definitions", sop_fields.get("4. Definitions", "")),
        ("5.0 References", sop_fields.get("5. References", "")),
        ("6.0 Procedure", sop_fields.get("6. Procedure", "")),
        ("6.1 Materials & Equipment", sop_fields.get("6.1 Materials & Equipment", "")),
        ("6.2 Stepwise Procedure", sop_fields.get("6.2 Stepwise Procedure", "")),
        ("7.0 Safety Precautions", sop_fields.get("7. Safety Precautions", "")),
        ("8.0 Training Requirements", sop_fields.get("8. Training Requirements", "")),
        ("9.0 Change Control", sop_fields.get("9. Change Control", "")),
        ("10.0 Records & Documentation", sop_fields.get("10. Records & Documentation", "")),
        ("11.0 Revision History", sop_fields.get("11. Revision History", "")),
    ]

    # Title paragraph (optional body text)
    if section_order[0][1]:
        _add_heading(doc, "Title", size=12, bold=True)
        _add_paragraphs(doc, _strip_markdown(section_order[0][1]))
        doc.add_paragraph()

    # Main sections up to 10.0 as normal text
    for heading, body in section_order[1:13]:  # 1.0 .. 10.0
        _add_heading(doc, heading, size=12, bold=True)
        _add_paragraphs(doc, _strip_markdown(body) or "—")
        doc.add_paragraph()

    # 8.0 heading already printed above; add the **Training Requirements table**
    _add_heading(doc, "Training Requirements", size=12, bold=True)
    _add_training_table(doc, rows=12)  # adjust row count as you like
    doc.add_paragraph()

    # 11.0 heading text followed by **Revision History table**
    _add_heading(doc, "Revision History", size=12, bold=True)
    _add_revision_history_table(
        doc,
        initial_version="v1.0",
        date_str=meta.get("effective_date", ""),
        initial_desc="Initial issue.",
        rows=10
    )
    doc.add_paragraph()

    # Footer
    _add_footer_with_form_and_page(doc, form_no=meta.get("form_no", "QA 01.04.02/14"))
    _add_page_border(doc, border_size=24, border_color="000000")


    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ---------------------------
# Summarised DOCX builder (bold headings like generated SOP)
# ---------------------------
RE_SUMMARY_HEAD = re.compile(
    r"^(Title|"
    r"\d\.0\s+Purpose|2\.0\s+Scope|3\.0\s+Responsibilities|4\.0\s+Definitions|5\.0\s+References|"
    r"6\.0\s+Procedure|6\.1\s+Materials\s*&\s*Equipment|6\.2\s+Stepwise\s+Procedure|"
    r"7\.0\s+Safety\s+Precautions|8\.0\s+Training\s+Requirements|9\.0\s+Change\s+Control|"
    r"10\.0\s+Records\s*&\s*Documentation|11\.0\s+Revision\s+History)\s*$",
    re.I
)

def _split_sop_style_text(txt: str) -> Dict[str, str]:
    """Split SOP-style summary text into sections by headings."""
    keys = [
        "Title",
        "1.0 Purpose","2.0 Scope","3.0 Responsibilities","4.0 Definitions","5.0 References",
        "6.0 Procedure","6.1 Materials & Equipment","6.2 Stepwise Procedure",
        "7.0 Safety Precautions","8.0 Training Requirements","9.0 Change Control",
        "10.0 Records & Documentation","11.0 Revision History",
    ]
    out = {k: "" for k in keys}
    cur = None
    for ln in (txt or "").splitlines():
        l = ln.strip()
        if not l:
            if cur:
                out[cur] += "\n"
            continue
        # detect heading
        for k in keys:
            if l.lower() == k.lower():
                cur = k
                continue
        if cur is None:
            # pre-title spill goes into Title body
            cur = "Title"
        else:
            out[cur] += (l + "\n")
    # strip trailing newlines
    return {k: v.strip() for k, v in out.items()}

def make_docx_summary_from_template(summary_text: str, meta: Dict[str, str]) -> bytes:
    """
    Build a summary .docx using the SAME layout as the generated SOP:
    - Header table with Title populated
    - 1.0 / 2.0 / … section headings
    - Cleaned body text (no ####, **, etc.)
    - Ending tables: Training Requirements and Revision History
    """
    # Fallback if python-docx isn't available
    if not DOCX_AVAILABLE:
        return (
            "STANDARD OPERATING PROCEDURE (SOP)\n"
            f"{meta.get('organization','')}\n\n"
            "SOP Summary\n\n"
            f"{summary_text}"
        ).encode("utf-8")

    doc = Document()

    # Header
    _add_heading(doc, "STANDARD OPERATING PROCEDURE (SOP)", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_heading(doc, meta.get("organization", ""), size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Header table (7x4) identical to SOP
    tbl = doc.add_table(rows=7, cols=4)
    tbl.style = "Table Grid"

    brand_cell = tbl.cell(0, 0).merge(tbl.cell(6, 0))
    brand_cell.text = "ASGS\nPharmaceuticals"

    # Parse the incoming summary into SOP-style sections (Title, 1.0, 2.0, …)
    sections = _split_sop_style_text(summary_text or "")
    # Title used in the header table (prefer parsed Title, fallback to meta title)
    header_title = _strip_markdown(sections.get("Title", "") or meta.get("title", ""))

    # Row 0
    tbl.cell(0, 1).text = "Department:"
    tbl.cell(0, 2).text = meta.get("department", "")
    tbl.cell(0, 3).text = f"SOP No.: {meta.get('sop_no','')}"
    # Row 1
    tbl.cell(1, 1).text = "Area:"
    tbl.cell(1, 2).text = meta.get("area", "")
    tbl.cell(1, 3).text = f"Effective Date: {meta.get('effective_date','')}"
    # Row 2
    tbl.cell(2, 1).text = "Title:"
    tbl.cell(2, 2).text = header_title
    tbl.cell(2, 3).text = f"Review Date: {meta.get('review_date','')}"
    # Row 3..6 scaffold
    tbl.cell(3, 1).text = "Prepared By"
    tbl.cell(3, 2).text = "Reviewed By"
    tbl.cell(3, 3).text = "Approved By"
    for i, label in enumerate(["Signature", "Name", "Date"], start=4):
        tbl.cell(i, 1).text = label
        tbl.cell(i, 2).text = label
        tbl.cell(i, 3).text = label

    doc.add_paragraph()

    # Body (match generated SOP formatting; clean markdown)
    ordered = [
        ("Title", sections.get("Title","")),
        ("1.0 Purpose", sections.get("1.0 Purpose","")),
        ("2.0 Scope", sections.get("2.0 Scope","")),
        ("3.0 Responsibilities", sections.get("3.0 Responsibilities","")),
        ("4.0 Definitions", sections.get("4.0 Definitions","")),
        ("5.0 References", sections.get("5.0 References","")),
        ("6.0 Procedure", sections.get("6.0 Procedure","")),
        ("6.1 Materials & Equipment", sections.get("6.1 Materials & Equipment","")),
        ("6.2 Stepwise Procedure", sections.get("6.2 Stepwise Procedure","")),
        ("7.0 Safety Precautions", sections.get("7.0 Safety Precautions","")),
        ("8.0 Training Requirements", sections.get("8.0 Training Requirements","")),
        ("9.0 Change Control", sections.get("9.0 Change Control","")),
        ("10.0 Records & Documentation", sections.get("10.0 Records & Documentation","")),
        ("11.0 Revision History", sections.get("11.0 Revision History","")),
    ]

    # Title section (optional body)
    if ordered[0][1]:
        _add_heading(doc, "Title", size=12, bold=True)
        _add_paragraphs(doc, _clean_section_body("Title", ordered[0][1]) or "—")
        doc.add_paragraph()

    # Sections 1.0 .. 10.0 as text
    for heading, body in ordered[1:13]:
        _add_heading(doc, heading, size=12, bold=True)
        _add_paragraphs(doc, _clean_section_body(heading, body) or "—")
        doc.add_paragraph()


    # Training Requirements table (same as generated SOP)
    _add_heading(doc, "Training Requirements", size=12, bold=True)
    _add_training_table(doc, rows=12)
    doc.add_paragraph()

    # Revision History table (same as generated SOP)
    _add_heading(doc, "Revision History", size=12, bold=True)
    _add_revision_history_table(
        doc,
        initial_version="v1.0",
        date_str=meta.get("effective_date", ""),
        initial_desc="Initial summary issue.",
        rows=10
    )
    doc.add_paragraph()

    # Footer consistency
    _add_footer_with_form_and_page(doc, form_no=meta.get("form_no", "QA 01.04.02/14"))
    _add_page_border(doc, border_size=24, border_color="000000")


    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ---------------------------
# === PDF-style Compliance & Quality Scoring (exact to screenshots) ===
# ---------------------------
import math

structure_sections = [
    "Title", "Purpose", "Scope", "Responsibilities", "Definitions",
    "References", "Procedure", "Safety", "Training", "Change Control",
    "Records", "Documentation", "Revision History", "Materials",
    "Equipment", "Stepwise"
]

regulatory_terms = [
    "21 CFR 210", "21 CFR 211", "21 CFR 11", "ICH", "EU GMP",
    "WHO GMP", "cGMP", "GMP", "PIC/S"
]

safety_terms = [
    "risk", "hazard", "PPE", "incident", "deviation", "CAPA",
    "spill", "lockout", "tagout", "exposure", "MSDS", "first aid"
]

compliance_terms = ["compliance", "audit", "regulatory", "requirement", "standard"]
traceability_terms = [r"\bVersion\s*\d+\b", r"\b\d{2}/\d{2}/\d{4}\b"]

weights_pdf = {
    "Structure": 10,
    "Regulatory Compliance": 15,
    "Safety Risk Coverage": 10,
    "Compliance Keywords Coverage": 10,
    "Traceability & Version Control": 10,
    "Technical Accuracy": 10,
    "Flesch Reading Ease": 5,
    "Gunning Fog Index": 5,
    "Avg Sentence Length": 5,
    "Long Sentences Count": 5,
    "Passive Voice Count": 5,
    "Bullet/Numbering Usage": 5,
    "ALL CAPS Section Count": 5,
}

ranges_pdf = {
    "Structure": (1, 0.5),
    "Regulatory Compliance": (1, 0.5),
    "Safety Risk Coverage": (1, 0.5),
    "Compliance Keywords Coverage": (1, 0.5),
    "Traceability & Version Control": (1, 0.5),
    "Technical Accuracy": (1, 0.5),
    "Flesch Reading Ease": (60, 40),                 # higher better
    "Gunning Fog Index": (12, 15, "low_better"),     # lower better
    "Avg Sentence Length": (20, 25, "low_better"),   # lower better
    "Long Sentences Count": (5, 10, "low_better"),   # lower better
    "Passive Voice Count": (5, 10, "low_better"),    # lower better
    "Bullet/Numbering Usage": (1, 0.5),              # normalized
    "ALL CAPS Section Count": (0, 2, "low_better"),  # lower better
}

try:
    import textstat
    _HAS_TEXTSTAT = True
except Exception:
    _HAS_TEXTSTAT = False

def _flesch(text: str) -> float:
    if _HAS_TEXTSTAT:
        try: return float(textstat.flesch_reading_ease(text))
        except Exception: pass
    words = re.findall(r'\w+', text)
    sents = [s for s in re.split(r'[.!?]+', text) if s.strip()]
    wpS = len(words) / max(1, len(sents))
    return max(0.0, 100 - (wpS - 14) * 5)

def _gunning_fog(text: str) -> float:
    if _HAS_TEXTSTAT:
        try: return float(textstat.gunning_fog(text))
        except Exception: pass
    words = re.findall(r'\w+', text)
    sents = [s for s in re.split(r'[.!?]+', text) if s.strip()]
    complex_w = sum(1 for w in words if len(w) >= 3)
    wpS = len(words) / max(1, len(sents))
    pct_complex = (complex_w / max(1, len(words))) * 100
    return 0.4 * (wpS + pct_complex)

def _normalize_fraction(count: int, total: int) -> float:
    return min(1.0, count / max(1, total)) if total else 0.0

def _score_from_range(name: str, value: float) -> float:
    spec = ranges_pdf[name]
    if len(spec) == 2:
        good, warn = float(spec[0]), float(spec[1])
        if good == 1.0 and warn == 0.5:
            return max(0.0, min(1.0, float(value)))
        if value >= good: return 1.0
        if value <= warn: return 0.0
        return (value - warn) / (good - warn)
    else:
        good, warn, mode = float(spec[0]), float(spec[1]), spec[2]
        if mode == "low_better":
            if value <= good: return 1.0
            if value >= warn: return 0.0
            return 1.0 - (value - good) / (warn - good)
        if value >= good: return 1.0
        if value <= warn: return 0.0
        return (value - warn) / (good - warn)

def _bullet_count(text: str) -> int:
    return len(re.findall(r"(?m)^\s*(?:[-*•]|\d+[.)])\s+", text))

def _passive_count(text: str) -> int:
    return len(re.findall(r'\b(?:is|are|was|were|be|been|being)\s+\w+ed\b', text, re.I))

def _all_caps_sections(text: str) -> int:
    return len(re.findall(r"(?m)^\s*[A-Z][A-Z ]{3,}\s*$", text))

def score_sop_pdf_metrics(text: str) -> Dict[str, float]:
    text_lc = text or ""
    found_sections = sum(1 for sec in structure_sections if re.search(rf"\b{re.escape(sec)}\b", text_lc, re.I))
    structure = _normalize_fraction(found_sections, len(structure_sections))

    reg_hits = sum(1 for term in regulatory_terms if re.search(re.escape(term), text_lc, re.I))
    safety_hits = sum(1 for term in safety_terms if re.search(re.escape(term), text_lc, re.I))
    comp_hits = sum(1 for term in compliance_terms if re.search(re.escape(term), text_lc, re.I))

    regulatory = _normalize_fraction(reg_hits, len(regulatory_terms))
    safety = _normalize_fraction(safety_hits, len(safety_terms))
    comp_cov = _normalize_fraction(comp_hits, len(compliance_terms))

    trace_hits = sum(len(re.findall(pat, text_lc)) for pat in traceability_terms)
    traceability = min(1.0, trace_hits / 2.0)

    technical = comp_cov  # per screenshot mapping

    fre = _flesch(text_lc)
    gfi = _gunning_fog(text_lc)

    sentences = [s for s in re.split(r'[.!?]+', text_lc) if s.strip()]
    words = re.findall(r'\w+', text_lc)
    avg_len = (len(words) / max(1, len(sentences))) if sentences else 0.0
    long_sent = sum(1 for s in sentences if len(s.split()) > 25)
    passive = _passive_count(text_lc)
    bullets = _bullet_count(text_lc)
    caps = _all_caps_sections(text_lc)

    raw = {
        "Structure": structure,
        "Regulatory Compliance": regulatory,
        "Safety Risk Coverage": safety,
        "Compliance Keywords Coverage": comp_cov,
        "Traceability & Version Control": traceability,
        "Technical Accuracy": technical,
        "Flesch Reading Ease": fre,
        "Gunning Fog Index": gfi,
        "Avg Sentence Length": avg_len,
        "Long Sentences Count": long_sent,
        "Passive Voice Count": passive,
        "Bullet/Numbering Usage": min(bullets / 5.0, 1.0),
        "ALL CAPS Section Count": caps,
    }
    return raw

def build_pdf_matrix(text: str) -> pd.DataFrame:
    raw = score_sop_pdf_metrics(text or "")
    norm = {k: _score_from_range(k, raw[k]) for k in raw.keys()}
    weighted = {k: round(norm[k] * weights_pdf[k], 4) for k in norm.keys()}
    total = round(sum(weighted.values()), 2)

    row = {}
    for k in raw.keys():
        row[f"{k} (raw)"] = raw[k]
    for k in norm.keys():
        row[f"{k} (0-1)"] = norm[k]
        row[f"{k} Weighted"] = weighted[k]
    row["Total Score (0-100)"] = total

    return pd.DataFrame([row])

# ---------------------------
# OpenAI helpers
# ---------------------------
_OPENAI_AVAILABLE_FLAG = _OPENAI_AVAILABLE

def _get_openai_key() -> Optional[str]:
    # Replace with your own secret management as needed
    return os.getenv("OPENAI_API_KEY")

@st.cache_resource(show_spinner=False)
def _get_openai_client():
    if not _OPENAI_AVAILABLE_FLAG:
        return None, "OpenAI SDK not installed. Run: pip install openai"
    api_key = _get_openai_key()
    if not api_key:
        return None, "Missing OpenAI API key (set st.secrets['OPENAI_API_KEY'] or env var OPENAI_API_KEY)"
    try:
        client = OpenAI(api_key=api_key)  # type: ignore
        return client, None
    except Exception:
        try:
            openai.api_key = api_key  # type: ignore
            return "legacy", None
        except Exception as e:
            return None, f"Failed to init OpenAI: {e}"

def _chat_completion(messages: List[Dict], model: str = "gpt-4o-mini", temperature: float = 0.3, max_tokens: int = 1800) -> Tuple[Optional[str], Optional[str]]:
    client, err = _get_openai_client()
    if err: return None, err
    if client == "legacy":
        try:
            resp = openai.ChatCompletion.create(model=model, messages=messages, temperature=temperature, max_tokens=max_tokens)  # type: ignore
            return resp["choices"][0]["message"]["content"].strip(), None
        except Exception as e:
            return None, str(e)
    try:
        resp = client.chat.completions.create(model=model, messages=messages, temperature=temperature, max_tokens=max_tokens)  # type: ignore
        return resp.choices[0].message.content.strip(), None
    except Exception as e:
        return None, str(e)


# ---------------------------
# Readability helpers + Optimization loop
# ---------------------------
def readability_scores(text: str) -> Tuple[float, float]:
    """Return (Flesch Reading Ease, Gunning Fog Index) using the
    same fallbacks already present in this file."""
    return _flesch(text or ""), _gunning_fog(text or "")

def _rewrite_for_readability(sop_text: str) -> Tuple[Optional[str], Optional[str]]:
    """LLM pass to improve readability while keeping SOP structure."""
    prompt = f"""
Rewrite the following SOP for better readability.
- Use short sentences (10–18 words)
- Prefer active voice
- Break procedures into bullets/numbers
- Avoid complex words unless required
- Keep all original headings and their order
- Target Flesch Reading Ease > 60 and Gunning Fog Index < 12

SOP:
{sop_text}
"""
    messages = [
        {"role": "system", "content": "You are an SOP readability optimizer."},
        {"role": "user", "content": prompt.strip()},
    ]
    return _chat_completion(messages, model="gpt-4o-mini", temperature=0.5, max_tokens=2200)

def generate_optimized_sop(topic: str, max_rounds: int = 3) -> Tuple[Optional[str], Optional[str], List[str]]:
    """
    Draft an SOP, then iteratively improve readability until targets are met
    or max_rounds is reached. Returns (sop_text, err, logs).
    """
    logs: List[str] = []
    sop, err = gpt_generate_full_sop(topic)
    if err or not sop:
        return None, (err or "Failed to draft SOP"), logs

    for i in range(max_rounds):
        fre, gfi = readability_scores(sop)
        logs.append(f"Round {i+1}: FRE={fre:.2f}, GFI={gfi:.2f}")
        if fre >= 60 and gfi <= 12:
            logs.append("✅ SOP meets readability targets.")
            return sop, None, logs

        logs.append("⚠️ Readability below threshold, rewriting…")
        rewritten, r_err = _rewrite_for_readability(sop)
        if r_err or not rewritten:
            logs.append(f"⚠️ Rewrite failed: {r_err or 'unknown error'}")
            break
        sop = rewritten

    logs.append("⚠️ Max optimization rounds reached.")
    return sop, None, logs

# ---------------------------
# Summary generation + readability optimization
# ---------------------------
def generate_optimized_summary(source_text: str, max_rounds: int = 2) -> Tuple[Optional[str], Optional[str], List[str]]:
    logs: List[str] = []
    if not (source_text or "").strip():
        return None, "No text extracted", logs

    summary, err = gpt_summarize_to_sop(source_text)
    if err or not summary:
        return None, (err or "Summarization failed"), logs

    for i in range(max_rounds):
        fre, gfi = readability_scores(summary)
        logs.append(f"Round {i+1} (summary): FRE={fre:.2f}, GFI={gfi:.2f}")
        if fre >= 60 and gfi <= 12:
            logs.append("✅ Summary meets readability targets.")
            return summary, None, logs

        logs.append("⚠️ Summary readability below threshold, rewriting…")
        rewritten, r_err = _rewrite_for_readability(summary)
        if r_err or not rewritten:
            logs.append(f"⚠️ Summary rewrite failed: {r_err or 'unknown error'}")
            break
        summary = rewritten

    logs.append("⚠️ Max optimization rounds reached for summary.")
    return summary, None, logs



# --- Abstractive, SOP-style summary (keeps 1.0/2.0 numbering) ---
def _force_x0_headings(text: str) -> str:
    lines = []
    for ln in (text or "").splitlines():
        m = re.match(r'^(\d+)(?!\.\d)(?:\.)?\s+(.*)$', ln.strip())
        if m:
            n, rest = m.groups()
            lines.append(f"{n}.0 {rest}".rstrip())
        else:
            lines.append(ln)
    return "\n".join(lines)

def _anti_copy_sanitize(source: str, summary: str, ngram: int = 5, max_overlap: float = 0.6) -> str:
    def ngrams(s: str, n: int) -> set:
        toks = re.findall(r"\w+|\S", s.lower())
        return set(tuple(toks[i:i+n]) for i in range(max(0, len(toks)-n+1)))
    src = ngrams(source or "", ngram)
    out_lines = []
    for ln in (summary or "").splitlines():
        if re.match(r'^(Title|(?:\d+(?:\.\d)?\s))', ln.strip(), flags=re.I):
            out_lines.append(ln); continue
        grams = ngrams(ln, ngram)
        overlap = (len(grams & src) / max(1, len(grams))) if grams else 0.0
        if overlap <= max_overlap:
            out_lines.append(ln)
        else:
            short = re.split(r'[.;:]', ln)[0].strip()
            if len(short.split()) >= 4:
                out_lines.append(short + ".")
    return "\n".join(out_lines).strip()

def gpt_summarize_to_sop(raw_text: str) -> Tuple[Optional[str], Optional[str]]:
    system = {
        "role": "system",
        "content": (
            "You are a pharma cGMP documentation specialist. Create an ABSTRACTIVE summary of a long SOP. "
            "Do NOT copy or quote sentences from the source. Rephrase everything in your own words. "
            "Target Flesch Reading Ease > 60, Gunning Fog Index < 12.\n"
            "Target 30–40% of the original length for each section.\n\n"
            "Output must use EXACTLY these headings in this order (no extras):\n"
            "Title\n1.0 Purpose\n2.0 Scope\n3.0 Responsibilities\n4.0 Definitions\n5.0 References\n"
            "6.0 Procedure\n6.1 Materials & Equipment\n6.2 Stepwise Procedure\n7.0 Safety Precautions\n"
            "8.0 Training Requirements\n9.0 Change Control\n10.0 Records & Documentation\n11.0 Revision History\n\n"
            "Formatting rules:\n"
            "- Plain text only (no markdown/bold).\n"
            "- Title: 1–2 lines.\n"
            "- 1.0–5.0, 7.0–10.0: 2–4 sentences each.\n"
            "- 6.0: 3–5 sentences overview.\n"
            "- 6.1: 3–6 short lines.\n"
            "- 6.2: numbered 1.–10. steps, one sentence each.\n"
            "- 11.0: one short line.\n"
            "Avoid any sequence of 12+ consecutive source words."
        )
    }
    user = {
        "role": "user",
        "content": "Summarize the following SOP content into the structure above. Compress aggressively and do NOT copy sentences.\n\n" + (raw_text or "")[:12000]
    }
    content, err = _chat_completion([system, user], model="gpt-4o-mini", temperature=0.15, max_tokens=1400)
    if err: return None, err
    if not content: return None, "Empty summary from model"
    out = _force_x0_headings(content.strip())
    out = _anti_copy_sanitize(raw_text or "", out, ngram=5, max_overlap=0.6)
    required = [
        "Title","1.0 Purpose","2.0 Scope","3.0 Responsibilities","4.0 Definitions","5.0 References",
        "6.0 Procedure","6.1 Materials & Equipment","6.2 Stepwise Procedure","7.0 Safety Precautions",
        "8.0 Training Requirements","9.0 Change Control","10.0 Records & Documentation","11.0 Revision History",
    ]
    for h in required:
        if h not in out:
            out += ("\n\n" + h + "\n—")
    return out.strip(), None

def gpt_generate_full_sop(topic: str) -> Tuple[Optional[str], Optional[str]]:
    sys = {"role": "system", "content": (
        "You write cGMP-compliant SOPs for pharma. Write formal, specific content.\n"
        "Target Flesch Reading Ease > 60, Gunning Fog Index < 12.\n"
        "Strictly follow the exact section order and headings below (no extra headings):\n"
        "Title\n"
        "1. Purpose\n"
        "2. Scope\n"
        "3. Responsibilities\n"
        "4. Definitions\n"
        "5. References\n"
        "6. Procedure\n"
        "6.1 Materials & Equipment\n"
        "6.2 Stepwise Procedure\n"
        "7. Safety Precautions\n"
        "8. Training Requirements\n"
        "9. Change Control\n"
        "10. Records & Documentation\n"
        "11. Revision History\n"
        "Do not include anything outside these headings."
    )}
    usr = {"role": "user", "content": (
        "Draft an SOP in Markdown for this topic: " + topic.strip() + "\n"
        "Target ~1,200–1,800 words. Use clear, numbered steps in 6.2."
    )}
    content, err = _chat_completion([sys, usr], model="gpt-4o-mini", temperature=0.3, max_tokens=3000)
    if err: return None, err
    return content, None

# ---------------------------
# Template parsing (generated SOP)
# ---------------------------
SECTION_KEYS = [
    "Title",
    "1. Purpose",
    "2. Scope",
    "3. Responsibilities",
    "4. Definitions",
    "5. References",
    "6. Procedure",
    "6.1 Materials & Equipment",
    "6.2 Stepwise Procedure",
    "7. Safety Precautions",
    "8. Training Requirements",
    "9. Change Control",
    "10. Records & Documentation",
    "11. Revision History",
]

import re

def _strip_markdown(text: str) -> str:
    if not text:
        return ""
    t = text
    t = re.sub(r"^#+\s*", "", t)              # remove leading ####
    t = re.sub(r"\*\*(.*?)\*\*", r"\1", t)    # **bold**
    t = re.sub(r"\*(.*?)\*", r"\1", t)        # *italic*
    t = re.sub(r"`(.*?)`", r"\1", t)          # `code`
    return t.strip()


def parse_sop_md(md: str) -> Dict[str, str]:
    lines = md.splitlines()
    content_map: Dict[str, List[str]] = {k: [] for k in SECTION_KEYS}
    current = None

    def _detect_header_and_inline(line: str) -> Tuple[Optional[str], Optional[str]]:
        raw = re.sub(r'^[#\s]+', '', line).strip()
        if not raw:
            return None, None
        norm = raw.replace('—', '-').replace('–', '-')

        for key in SECTION_KEYS:
            if norm.lower().startswith(key.lower()):
                rest = norm[len(key):].lstrip(": -")
                return key, _strip_markdown(rest) if rest else None
        return None, None

    for ln in lines:
        hdr, inline = _detect_header_and_inline(ln)
        if hdr:
            current = hdr
            if inline:
                content_map[current].append(inline)
            continue
        if current:
            content_map[current].append(_strip_markdown(ln))

    # Final cleanup
    return {k: "\n".join(v).strip() for k, v in content_map.items()}



# ---------------------------
# Session state defaults
# ---------------------------
for key, default in [
    ("audit", []), ("summary", ""), ("summary_pdf", None), ("summary_docx", None),
    ("chat_history", []), ("generated_sop_md", ""), ("compliance_df", None),
    ("compliance_total", None), ("generated_sop_pdf", None), ("generated_sop_docx", None),
    ("jump_to_uploads", False),
    ("summary_matrix_df", None), ("summary_quality_score", None),
    ("last_file_id", None),
]:
    if key not in st.session_state: st.session_state[key] = default

# ---------------------------
# Header
# ---------------------------
st.markdown("<h1 style='text-align:center; letter-spacing:1px;'>REGDOCGPT</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align:center; color:#5b6770; font-size:18px;'>Upload to Summarize • Chat to Generate SOPs</p>", unsafe_allow_html=True)
st.markdown("---")

# ---------------------------
# Sidebar (API status + whoami + SOP metadata inputs)
# ---------------------------
with st.sidebar:
    st.subheader("OpenAI status")
    key_present = bool(_get_openai_key())
    st.write("API key:", "✅ found" if key_present else "❌ missing")
    st.caption("Store the key in st.secrets['OPENAI_API_KEY'] or the OPENAI_API_KEY env var.")
    st.markdown("---")
    st.write(f"**Logged in:** {acct.get('username','user')}")
    st.write(f"**Role:** {acct.get('role','user').title()}")
    st.write(f"**User ID:** {acct.get('id','—')}")
    st.markdown("---")
    st.subheader("SOP Metadata")
    meta_department = st.text_input("Department", value="", placeholder="Quality Assurance")
    meta_area = st.text_input("Area", value="", placeholder="Manufacturing")
    meta_sopno = st.text_input("SOP No.", value="", placeholder="SOP-XXX")
    meta_effective = st.date_input("Effective Date", value=datetime.now()).strftime("%Y-%m-%d")
    meta_review = st.text_input("Review Date (optional)", value="")
    meta_org = st.text_input("Organization", value=acct.get("org", "ASGS Pharmaceuticals"))
    meta_form_no = st.text_input("Form No.", value="QA 01.04.02/14")

# ---------------------------
# Layout
# ---------------------------
left, right = st.columns([1, 1])

# ---- LEFT: Upload & Summary ----
with left:
    if st.session_state.get("jump_to_uploads"):
        st.info("📂 My documents / uploads")
        st.session_state["jump_to_uploads"] = False

    st.markdown("### Upload SOP (auto-summarize)")
    uploaded = st.file_uploader("Browse files (.txt/.pdf/.docx)", type=["txt", "pdf", "docx"], label_visibility="collapsed")

    if uploaded is not None:
        # ---- Prevent re-summarising on every rerun (use name+size signature)
        content_bytes = uploaded.read()
        file_sig = (uploaded.name, len(content_bytes))
        need_process = (st.session_state.last_file_id != file_sig)

        if need_process:
            ext = os.path.splitext(uploaded.name)[1].lower()
            extracted_text = ""
            try:
                if ext == ".txt":
                    extracted_text = content_bytes.decode("utf-8", errors="ignore")
                elif ext == ".pdf":
                    if fitz is None:
                        st.warning("PyMuPDF not installed; cannot parse PDF.")
                    else:
                        doc = fitz.open(stream=content_bytes, filetype="pdf")
                        extracted_text = "\n".join(page.get_text() for page in doc)
                elif ext == ".docx":
                    if docx2txt is None:
                        st.warning("docx2txt not installed; cannot parse .docx.")
                    else:
                        tmp = "_tmp_upload.docx"
                        with open(tmp, "wb") as f:
                            f.write(content_bytes)
                        try:
                            extracted_text = docx2txt.process(tmp)
                        finally:
                            try: os.remove(tmp)
                            except Exception: pass
            except Exception as e:
                st.error(f"Failed to extract text: {e}")
                extracted_text = ""

            # === SOP-style Summary (abstractive) ===
            with st.spinner("Summarizing SOP…"):
                sop_style_summary, err = gpt_summarize_to_sop(extracted_text) if extracted_text.strip() else (None, "No text extracted")
                if err or not sop_style_summary:
                    st.info(f"Falling back to heuristic: {err or 'no content'}")
                    sop_style_summary = summarize_heuristic(extracted_text)

            st.session_state.summary = sop_style_summary

            # Build summary DOCX (header/table same as SOPs; bold headings)
            meta_summary = {
                "organization": meta_org or "ASGS Pharmaceuticals",
                "department": meta_department,
                "sop_no": meta_sopno,
                "area": meta_area,
                "effective_date": meta_effective,
                "review_date": meta_review,
                "title": os.path.splitext(uploaded.name)[0] if uploaded is not None else "SOP Summary",
                "form_no": meta_form_no or "QA 01.04.02/14",
            }
            st.session_state.summary_docx = make_docx_summary_from_template(st.session_state.summary or "", meta_summary)

            # We no longer generate a Summary PDF
            st.session_state.summary_pdf = None

            # === Compliance Matrix + Quality Score for SUMMARY (PDF-style) ===
            matrix_df = build_pdf_matrix(st.session_state.summary or "")
            final_score = float(matrix_df.loc[0, "Total Score (0-100)"])
            st.session_state.summary_matrix_df = matrix_df
            st.session_state.summary_quality_score = final_score

            add_audit("User", "SOP uploaded", uploaded.name)
            add_audit("System", "SOP summarized", uploaded.name)

            # Remember processed file signature
            st.session_state.last_file_id = file_sig

    st.markdown("### Summary")
    st.text_area("Preview", value=st.session_state.summary or "", height=260)

    # Show SUMMARY quality score & matrix
    if st.session_state.summary_matrix_df is not None:
        st.metric("Summary Quality and Compliance Score", f"{st.session_state.summary_quality_score}/100")
        st.dataframe(st.session_state.summary_matrix_df, use_container_width=True)
        st.download_button(
            "Download Summary Compliance Matrix (CSV)",
            data=st.session_state.summary_matrix_df.to_csv(index=False).encode("utf-8"),
            file_name="summary_compliance_matrix.csv",
            mime="text/csv",
            use_container_width=True,
            on_click=lambda: add_audit("Assistant", "Download", "summary_compliance_matrix.csv"),
        )

    st.markdown("#### Download summary")
    disabled = not bool(st.session_state.summary_docx)
    st.download_button(
        "Download Summary (Word)",
        data=st.session_state.summary_docx or b"",
        file_name="sop_summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        disabled=disabled,
        use_container_width=True,
        on_click=lambda: add_audit("User", "Download", "sop_summary.docx"),
    )

# ---- RIGHT: Chatbot & Generated SOP ----
with right:
    st.markdown("### Generate SOP")
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])

    user_prompt = st.chat_input("Ask about the SOP, or say: generate sop for rapid mixer granulator…")
    if user_prompt:
        st.session_state.chat_history.append({"role": "user", "content": user_prompt})
        with st.chat_message("user"): st.write(user_prompt)

        text = user_prompt.strip()
        lower = text.lower()

        patterns = [
            r'\b(?:generate|create|draft|write|make|prepare|build|produce)\b.*?\b(?:sop|standard operating procedure)\b\s*(?:for|on|about|:)?\s*(.+)$',
            r'^(?:sop|standard operating procedure)\s*(?:for|on|about|:)\s*(.+)$',
        ]

        topic = None
        for pat in patterns:
            m = re.search(pat, lower)
            if m:
                topic = text[m.start(1):m.end(1)].strip(" .-–—")
                break

        if topic:
            with st.spinner("Drafting SOP..."):
              sop_md, err, rounds_log = generate_optimized_sop(topic, max_rounds=3)


            if err or not sop_md:
                reply = f"Generation failed: {err or 'no content'}"
                st.session_state.chat_history.append({"role": "assistant", "content": reply})
                with st.chat_message("assistant"): st.write(reply)
            else:
                # Save content
                st.session_state.generated_sop_md = sop_md

                # Compute Compliance Matrix + Quality Score for GENERATED SOP
                gen_matrix_df = build_pdf_matrix(sop_md)
                gen_final_score = float(gen_matrix_df.loc[0, "Total Score (0-100)"])
                st.session_state.compliance_df = gen_matrix_df
                st.session_state.compliance_total = gen_final_score

                # Build the final Word & PDF
                sop_fields = parse_sop_md(sop_md)
                meta = {
                    "organization": meta_org or "ASGS Pharmaceuticals",
                    "department": meta_department,
                    "sop_no": meta_sopno,
                    "area": meta_area,
                    "effective_date": meta_effective,
                    "review_date": meta_review,
                    "title": (sop_fields.get("Title") or topic or "Standard Operating Procedure"),
                    "form_no": meta_form_no or "QA 01.04.02/14",
                }
                st.session_state.generated_sop_docx = make_docx_from_template(sop_fields, meta)
                st.session_state.generated_sop_pdf  = make_pdf_from_template(sop_fields, meta)

                add_audit("Assistant", "Generated SOP", topic[:120])
                with st.chat_message("assistant"):
                    st.write(f"Drafted the SOP for **{topic}** and calculated a quality/compliance score. Scroll to **Generated SOP & Downloads**.")
        else:
            if not st.session_state.summary:
                reply = ("I can generate a new SOP too. Try phrases like:\n"
                         "• generate sop for rapid mixer granulator\n"
                         "• write standard operating procedure about equipment cleaning")
            else:
                s = st.session_state.summary
                if any(k in lower for k in ["summary", "overview", "tl;dr", "what is this", "what does it say"]):
                    reply = f"Here’s the current summary of the SOP:\n\n{s}"
                elif any(k in lower for k in ["next step", "action", "what should", "checklist"]):
                    reply = ("Based on the SOP summary, suggested next steps:\n"
                             "1) Validate responsible roles & approvals\n"
                             "2) Confirm training requirements\n"
                             "3) Verify effective/retirement dates\n"
                             "4) Ensure change control links and related SOPs are referenced\n"
                             "5) Publish and notify impacted teams")
                elif "risk" in lower:
                    reply = ("From the summary, check for:\n"
                             "- Deviations and CAPA handling\n"
                             "- Data integrity controls (ALCOA+)\n"
                             "- Change control and impact assessment\n"
                             "- Batch record review steps\n"
                             "If any are absent, flag as a potential gap.")
                else:
                    reply = ("Here’s what I can infer from the current SOP summary:\n\n"
                             f"{s}\n\n"
                             "You can also ask me to generate a new SOP, e.g., 'generate sop for rapid mixer granulator'.")
            st.session_state.chat_history.append({"role": "assistant", "content": reply})
            add_audit("Assistant", "Replied", reply[:180] + ("…" if len(reply) > 180 else ""))
            with st.chat_message("assistant"): st.write(reply)

    if st.session_state.generated_sop_md:
        st.markdown("---"); st.markdown("### Generated SOP & Downloads")
        st.text_area("SOP (Markdown preview)", value=st.session_state.generated_sop_md, height=240)
        if st.session_state.compliance_df is not None:
            st.metric("Generated SOP Quality and Compliance Score", f"{st.session_state.compliance_total}/100")
            st.dataframe(st.session_state.compliance_df, use_container_width=True)
            st.download_button("Download Generated SOP Compliance Matrix (CSV)",
                               data=st.session_state.compliance_df.to_csv(index=False).encode("utf-8"),
                               file_name="generated_sop_compliance_matrix.csv", mime="text/csv")

        st.download_button("Download SOP (Word)",
                           data=st.session_state.generated_sop_docx or b"",
                           file_name=f"{meta_sopno or 'generated_sop'}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
