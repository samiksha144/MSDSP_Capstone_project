
import io
import time
from datetime import datetime
import streamlit as st
import pandas as pd

from backend import Services

# ---------- Page config ----------
st.set_page_config(page_title="REGDOCGPT", page_icon="📄", layout="wide")

# Initialize services (cache across reruns)
@st.cache_resource(show_spinner=False)
def get_services():
    return Services()

svc = get_services()

# ---------- Utilities ----------
def now_iso():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def add_audit(actor: str, event: str, detail: str = ""):
    st.session_state.audit.insert(
        0,
        {"Timestamp": now_iso(), "Actor": actor, "Event": event, "Detail": detail},
    )

# ---------- Session state ----------
for k, v in [
    ("audit", []),
    ("summary", ""),
    ("uploaded_name", None),
    ("pdf_bytes", None),
    ("docx_bytes", None),
    ("chat_history", []),
    ("raw_text", ""),
]:
    if k not in st.session_state:
        st.session_state[k] = v

# ---------- Header ----------
st.markdown("<h1 style='text-align:center; letter-spacing:1px;'>REGDOCGPT</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align:center; color:#5b6770; font-size:18px;'>Generative AI Assistant for Regulatory Compliance in Pharma</p>", unsafe_allow_html=True)
st.markdown("---")

# ---------- Layout ----------
left, right = st.columns([1, 1])

with left:
    st.markdown("### Upload SOP")
    uploaded = st.file_uploader("Browse files", type=None, label_visibility="collapsed")

    if uploaded is not None:
        content = uploaded.read()
        st.session_state.uploaded_name = uploaded.name
        add_audit("User", "SOP uploaded", uploaded.name)

        with st.spinner("Extracting & summarizing…"):
            time.sleep(0.4)
            res = svc.summarize(content, filename=uploaded.name)
            st.session_state.summary = res.text
            st.session_state.raw_text = content.decode("utf-8", errors="ignore") if isinstance(content, (bytes, bytearray)) else str(content)
            # Simple outputs (optional): generate quick PDF/DOCX stubs from summary
            try:
                from reportlab.lib.pagesizes import LETTER
                from reportlab.pdfgen import canvas
                bio = io.BytesIO()
                c = canvas.Canvas(bio, pagesize=LETTER)
                width, height = LETTER
                textobj = c.beginText(72, height - 72)
                textobj.setFont("Helvetica", 12)
                textobj.textLine("SOP Summary")
                textobj.textLine("")
                import textwrap as _tw
                for line in _tw.wrap(st.session_state.summary or "", width=88):
                    textobj.textLine(line)
                c.drawText(textobj); c.showPage(); c.save()
                st.session_state.pdf_bytes = bio.getvalue()
            except Exception:
                st.session_state.pdf_bytes = (st.session_state.summary or "").encode("utf-8")

            try:
                from docx import Document
                doc = Document()
                doc.add_heading("SOP Summary", level=1)
                doc.add_paragraph(st.session_state.summary or "")
                b = io.BytesIO(); doc.save(b)
                st.session_state.docx_bytes = b.getvalue()
            except Exception:
                st.session_state.docx_bytes = (st.session_state.summary or "").encode("utf-8")

        add_audit("System", "SOP summarized", f"Source: {uploaded.name}")

with right:
    st.markdown("### Workspace")
    tab_summary, tab_chat = st.tabs(["Summary", "Chatbot"])

    with tab_summary:
        st.text_area(
            "Summary of uploaded SOP",
            value=st.session_state.summary or "",
            height=220,
            label_visibility="collapsed",
        )

    with tab_chat:
        colA, colB, colC = st.columns([1, 1, 1])
        with colA:
            if st.button("Clear chat", use_container_width=True):
                st.session_state.chat_history = []
                add_audit("User", "Chat cleared")
        with colB:
            generate_from_summary = st.button("Generate SOP (from summary)", use_container_width=True)
        with colC:
            st.caption("Tip: Type \"generate sop: <brief>\"")

        # Render history
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        user_prompt = st.chat_input("Ask about the SOP or type: generate sop: Cleaning bioreactor vessels…")
        if user_prompt:
            st.session_state.chat_history.append({"role": "user", "content": user_prompt})
            add_audit("User", "Asked", user_prompt)
            with st.chat_message("user"):
                st.write(user_prompt)

            q = user_prompt.strip()
            ql = q.lower()

            # --- SOP generation intent -------------------------------------------------
            sop_requested = False
            brief = None
            for trigger in ["generate sop:", "create sop:", "draft sop:", "make sop:"]:
                if ql.startswith(trigger):
                    sop_requested = True
                    brief = q[len(trigger):].strip()
                    break
            if not sop_requested and generate_from_summary:
                sop_requested = True
                brief = st.session_state.summary or st.session_state.raw_text or "Create a general SOP from the summary."

            if sop_requested:
                with st.spinner("Generating SOP…"):
                    sop = svc.generate_sop(brief or q)
                add_audit("Assistant", "SOP generated", (brief or q)[:140])
                st.session_state.chat_history.append({"role": "assistant", "content": "Here is the generated SOP (JSON):"})
                with st.chat_message("assistant"):
                    st.json(sop)
            else:
                # --- QA behavior backed by summary text --------------------------------
                s = (st.session_state.summary or "").strip()
                if not s:
                    reply = "I don’t see an SOP uploaded yet. Please upload a document, then ask me about its summary or type: generate sop: ..."
                else:
                    if any(k in ql for k in ["summary", "overview", "tl;dr", "what is this", "what does it say"]):
                        reply = f"Here’s the current summary of the SOP:\n\n{s}"
                    elif any(k in ql for k in ["next step", "action", "what should", "checklist"]):
                        reply = (
                            "Based on the SOP summary, suggested next steps:\n"
                            "1) Validate responsible roles & approvals\n"
                            "2) Confirm training requirements\n"
                            "3) Verify effective/retirement dates\n"
                            "4) Ensure change control links and related SOPs are referenced\n"
                            "5) Publish and notify impacted teams"
                        )
                    elif "risk" in ql:
                        reply = (
                            "From the summary, check for:\n"
                            "- Deviations and CAPA handling\n"
                            "- Data integrity controls (ALCOA+)\n"
                            "- Change control and impact assessment\n"
                            "- Batch record review steps\n"
                            "If any are absent, flag as a potential gap."
                        )
                    else:
                        reply = (
                            "Here’s what I can infer from the current SOP summary:\n\n"
                            f"{s}\n\n"
                            "You can also ask me to *generate* a structured SOP: `generate sop: <brief>`"
                        )
                st.session_state.chat_history.append({"role": "assistant", "content": reply})
                add_audit("Assistant", "Replied", reply[:180] + ("…" if len(reply) > 180 else ""))
                with st.chat_message("assistant"):
                    st.write(reply)

# ---------- Outputs ----------
st.markdown("---")
st.markdown("### Outputs")
c1, c2 = st.columns([3, 1])
with c1:
    st.caption("Generate files from the current summary.")
with c2:
    disabled = not bool(st.session_state.summary)
    def _on_pdf(): add_audit("User", "Download", "sop_summary.pdf")
    def _on_docx(): add_audit("User", "Download", "sop_summary.docx")

    st.download_button(
        label="Download PDF",
        data=st.session_state.pdf_bytes or b"",
        file_name="sop_summary.pdf",
        mime="application/pdf",
        disabled=disabled,
        use_container_width=True,
        on_click=_on_pdf,
    )
    st.download_button(
        label="Download Word",
        data=st.session_state.docx_bytes or b"",
        file_name="sop_summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        disabled=disabled,
        use_container_width=True,
        on_click=_on_docx,
    )

# ---------- Audit Trail ----------
st.markdown("---")
st.markdown("### Audit Trail")
audit_df = pd.DataFrame(st.session_state.audit)
if audit_df.empty:
    st.caption("No audit events yet.")
else:
    fcol1, fcol2, fcol3 = st.columns([1, 1, 2])
    with fcol1:
        actor_filter = st.multiselect("Actor", options=sorted(audit_df["Actor"].unique()), default=list(sorted(audit_df["Actor"].unique())))
    with fcol2:
        event_filter = st.multiselect("Event", options=sorted(audit_df["Event"].unique()), default=list(sorted(audit_df["Event"].unique())))
    with fcol3:
        search = st.text_input("Search detail", value="", placeholder="Contains…")

    filtered = audit_df[audit_df["Actor"].isin(actor_filter) & audit_df["Event"].isin(event_filter)]
    if search.strip():
        s = search.strip().lower()
        filtered = filtered[filtered["Detail"].str.lower().str.contains(s, na=False)]

    st.dataframe(filtered, hide_index=True, use_container_width=True)

    csv = filtered.to_csv(index=False).encode("utf-8")
    st.download_button("Export audit CSV", data=csv, file_name="audit_trail.csv", mime="text/csv", use_container_width=False)

st.markdown("<div style='text-align:right; color:#98a2b3; font-size:12px;'></div>", unsafe_allow_html=True)
