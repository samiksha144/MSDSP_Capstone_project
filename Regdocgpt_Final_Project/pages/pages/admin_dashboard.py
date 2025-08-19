import io
import os
import re
import json
import textwrap
from datetime import datetime
from typing import List, Tuple, Dict, Optional

import pandas as pd
import streamlit as st

# === DB audits integration (writes + reads for admin) ===
from db_repo import (
    ensure_audit_schema,
    add_audit as db_add_audit,
    get_audits as db_get_audits,
    DBError,
)

# ---------------------------
# Page config
# ---------------------------
st.set_page_config(page_title="REGDOCGPT — Admin", page_icon="🛡️", layout="wide")

# ---------------------------
# Page background
# ---------------------------
def set_grey_bg(color: str = "#000000"):
    st.markdown(
        f"""
        <style>
          [data-testid="stAppViewContainer"] {{
            background: {color} !important;
          }}
          [data-testid="stHeader"] {{
            background: {color} !important;
            z-index: 2 !important;
            border-bottom: 1px solid #e5e7eb;
          }}
          .block-container {{
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-radius: 16px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.05);
            padding-top: 1.2rem;
            padding-bottom: 1.2rem;
          }}
        </style>
        """,
        unsafe_allow_html=True,
    )

set_grey_bg("#000000")

# Hide Streamlit's top-right Deploy/status toolbar
def _hide_deploy_button():
    st.markdown(
        """
        <style>
        div[data-testid="stToolbar"] { display: none !important; }
        div[data-testid="stStatusWidget"] { display: none !important; }
        </style>
        """,
        unsafe_allow_html=True,
    )

_hide_deploy_button()

# ======== Require login & ensure admin ========
acct = st.session_state.get("account")
if not acct:
    st.switch_page("pages/login.py")
if (acct.get("role") or "").lower() != "admin":
    st.error("Admin access only.")
    st.stop()

# --- Profile menu (kept identical) ---
action = st.query_params.get("action")
if isinstance(action, list):
    action = action[0] if action else None

def _clear_action_param():
    try:
        del st.query_params["action"]
    except Exception:
        pass

if action == "logout":
    st.session_state.pop("account", None)
    _clear_action_param()
    st.switch_page("pages/login.py")


elif action == "edit_profile":
    with st.modal("Edit profile"):
        name = st.text_input("Name", value=acct.get("username", ""))
        email = st.text_input("Email", value=acct.get("email", ""))
        org = st.text_input("Organization", value=acct.get("org", ""))

        c1, c2 = st.columns(2)
        save = c1.button("Save")
        cancel = c2.button("Cancel", type="secondary")

        if save:
            # TODO: wire to SQL proc
            acct.update({"username": name, "email": email, "org": org})
            st.session_state["account"] = acct
            st.success("Profile updated (local). We'll wire DB next.")
            _clear_action_param()
            st.rerun()
        if cancel:
            _clear_action_param()
            st.rerun()

elif action == "change_password":
    with st.modal("Change password"):
        cur = st.text_input("Current password", type="password")
        new = st.text_input("New password", type="password")
        rep = st.text_input("Confirm new password", type="password")

        c1, c2 = st.columns(2)
        update = c1.button("Update")
        cancel = c2.button("Cancel", type="secondary")

        if update:
            if not new or new != rep:
                st.error("New passwords do not match.")
            else:
                # TODO: wire to SQL proc
                st.success("Password updated (placeholder). We'll wire DB next.")
                _clear_action_param()
                st.rerun()
        if cancel:
            _clear_action_param()
            st.rerun()

elif action == "my_docs":
    st.session_state["jump_to_uploads"] = True
    _clear_action_param()
    st.rerun()

def show_profile_menu():
    name = acct.get("username", "admin")
    uid  = acct.get("id", "—")

    html = f"""
<style>
[data-testid="stHeader"] {{ background: transparent; z-index: 0 !important; }}
.profile-menu {{ position: fixed; top: 10px; right: 16px; z-index: 999999 !important;
  font-family: ui-sans-serif,-apple-system,Segoe UI,Roboto,Helvetica,Arial; }}
.profile-menu details {{ position: relative; }}
.profile-menu summary {{ list-style: none; outline: none; cursor: pointer; display: flex; align-items: center; gap: 10px;
  padding: 6px 12px; background: rgba(255,255,255,0.98); border: 1px solid #e5e7eb; border-radius: 9999px;
  box-shadow: 0 8px 22px rgba(0,0,0,.10); user-select: none; }}
.profile-menu summary::-webkit-details-marker {{ display:none; }}
.profile-menu .name {{ font-weight: 700; color:#111827; line-height: 1; }}
.profile-menu .meta {{ font-size: 12px; color:#6b7280; margin-top:-2px; }}
.profile-menu .menu {{ position: absolute; right: 0; top: 100%; margin-top: 8px; min-width: 240px; background: #fff;
  border: 1px solid #e5e7eb; border-radius: 12px; box-shadow: 0 16px 40px rgba(0,0,0,.18); padding: 6px; }}
.profile-menu a {{ display: block; padding: 10px 12px; border-radius: 8px; text-decoration: none; color: #111827; }}
.profile-menu a:hover {{ background: #f3f4f6; }}
</style>
<div class="profile-menu">
<details>
<summary title="Open profile menu">
  <span style="font-size:18px;">🛡️</span>
  <div>
    <div class="name">{name}</div>
    <div class="meta">ID: {uid}</div>
  </div>
  <span style="margin-left:6px;">▾</span>
</summary>
<div class="menu">
  <a href="?action=logout" target="_self">Sign out</a>
</div>
</details>
</div>
"""
    st.markdown(html, unsafe_allow_html=True)

show_profile_menu()

# ---------------------------
# Optional deps
# ---------------------------
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None
try:
    import docx2txt
except Exception:
    docx2txt = None
try:
    import jinja2
except Exception:
    jinja2 = None

# ---- OpenAI (new SDK preferred; legacy supported) ----
_OPENAI_AVAILABLE = False
try:
    from openai import OpenAI  # >=1.0
    _OPENAI_AVAILABLE = True
except Exception:
    try:
        import openai  # legacy 0.x
        _OPENAI_AVAILABLE = True
    except Exception:
        _OPENAI_AVAILABLE = False

# ---------------------------
# DB bootstrap (audits table + procs) — run once
# ---------------------------
@st.cache_resource(show_spinner=False)
def _bootstrap_audit_schema() -> Optional[str]:
    try:
        ensure_audit_schema()
        return None
    except Exception as e:
        return f"{e}"

_bootstrap_err = _bootstrap_audit_schema()
if _bootstrap_err:
    st.warning(f"Audit DB bootstrap warning: {_bootstrap_err}")

# ---------------------------
# Utilities
# ---------------------------
def now_iso() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def _safe_int(x) -> Optional[int]:
    try:
        return int(x)
    except Exception:
        return None

# ---- add_audit wrapper (v2→v1 fallback) ----
def add_audit(actor: str, event: str, detail: str = "") -> None:
    acct_obj = st.session_state.get("account") or {}
    role = (acct_obj.get("role") or "user").lower()
    uid = _safe_int(acct_obj.get("id"))
    admin_id = uid if role == "admin" else None
    user_id = uid if role != "admin" else None

    try:
        db_add_audit(
            actor=actor,
            actor_role=role if role in ("user", "admin", "system", "assistant") else None,
            admin_id=admin_id,
            user_id=user_id,
            event=event,
            detail=detail,
        )
        return
    except TypeError:
        try:
            db_add_audit(
                actor=actor,
                user_id=user_id if user_id is not None else admin_id,
                event=event,
                detail=detail,
            )
            return
        except Exception as e:
            _append_audit_memory(actor, uid, event, f"(DB v1 fail) {detail} | {e}")
    except Exception as e:
        _append_audit_memory(actor, uid, event, f"(DB v2 fail) {detail} | {e}")

def _append_audit_memory(actor: str, uid: Optional[int], event: str, detail: str) -> None:
    if "audit" not in st.session_state:
        st.session_state.audit = []
    st.session_state.audit.insert(0, {
        "Timestamp": now_iso(),
        "Actor": actor,
        "UserID": str(uid) if uid is not None else "—",
        "Event": event,
        "Detail": detail,
    })

def summarize_heuristic(text: str) -> str:
    cleaned = " ".join((text or "").split())
    if not cleaned:
        return "The SOP document outlines the procedures for regulatory compliance and quality management."
    return (cleaned[:1200] + "…") if len(cleaned) > 1200 else cleaned

# ---- Simple multi-page PDF maker ----
def make_pdf_from_text(title: str, body: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import LETTER  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception:
        return f"{title}\n\n{body}".encode("utf-8")
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=LETTER)
    width, height = LETTER
    margin_x, margin_y = 72, 72
    max_width_chars = 95
    y = height - margin_y
    textobj = c.beginText(margin_x, y)
    textobj.setFont("Helvetica-Bold", 14)
    textobj.textLine(title)
    textobj.moveCursor(0, 16)
    textobj.setFont("Helvetica", 11)
    def flush_page():
        nonlocal textobj
        c.drawText(textobj); c.showPage()
        textobj = c.beginText(margin_x, height - margin_y); textobj.setFont("Helvetica", 11)
    for paragraph in body.split("\n"):
        lines = textwrap.wrap(paragraph, width=max_width_chars) if paragraph.strip() else [""]
        for line in lines:
            if textobj.getY() <= margin_y: flush_page()
            textobj.textLine(line)
        if textobj.getY() <= margin_y: flush_page()
        textobj.textLine("")
    c.drawText(textobj); c.showPage(); c.save()
    return bio.getvalue()

def make_docx_from_text(title: str, body: str) -> bytes:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return f"{title}\n\n{body}".encode("utf-8")
    doc = Document()
    doc.add_heading(title, level=1)
    for para in body.split("\n"):
        doc.add_paragraph(para)
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

# ---------------------------
# Compliance scoring
# ---------------------------
RE_HEADER = re.compile(r'^(#{1,3})\s*(.+?)\s*$', re.M)
RE_NUM_STEP = re.compile(r'(?m)^\s*(?:\d+[\).\]]|step\s*\d+[\).\]]?)\s+\S+')
RE_ROLES = re.compile(r'\b(QA|QC|Quality|Production|Operator|Supervisor|Engineering|Maintenance|Validation|EHS|Training|Manufacturing)\b', re.I)
RE_DOC_IDS = re.compile(r'\b(SOP|WI|FRM|LOG|REC|BMR|BRR)[-_ ]?\d{2,}[A-Z0-9-]*\b', re.I)
ALCOA_TERMS = [r'\bALCOA\+?\b', r'\battributable\b', r'\blegible\b', r'\bcontemporaneous\b', r'\boriginal\b', r'\baccurate\b', r'\bcomplete\b', r'\bconsistent\b', r'\benduring\b', r'\bavailable\b']
REFERENCE_TERMS = [r'\b21\s*CFR\s*210\b', r'\b21\s*CFR\s*211\b', r'\b21\s*CFR\s*11\b', r'\bICH\s*Q[0-9]+\b', r'\bEU\s*GMP\b', r'\bWHO\s*GMP\b', r'\bPIC/S\b']
WEIGHTS = {"Purpose": 0.10,"Scope": 0.10,"Responsibilities": 0.12,"Procedure": 0.25,"Change Control": 0.12,"Records / Documentation": 0.13,"Data Integrity (ALCOA)": 0.13,"Regulatory References": 0.05}

def _split_sections(md: str):
    md = md or ""; sections = {}; matches = list(RE_HEADER.finditer(md))
    if not matches: return {"_full": md}
    for i, m in enumerate(matches):
        start = m.end(); end = matches[i+1].start() if i+1 < len(matches) else len(md)
        header = m.group(2).strip().lower(); sections[header] = md[start:end].strip()
    return sections

def _len_score(txt: str, min_chars: int, good_chars: int) -> float:
    n = len((txt or "").strip())
    if n <= min_chars: return 0.0
    if n >= good_chars: return 1.0
    return (n - min_chars) / max(1, (good_chars - min_chars))

def evaluate_sop_compliance(md_text: str) -> Tuple[pd.DataFrame, float]:
    sec = _split_sections(md_text)
    def get_sec(*aliases):
        for a in aliases:
            for k, v in sec.items():
                if a in k: return v
        return ""
    s_purpose = get_sec("purpose"); score_purpose = _len_score(s_purpose, 200, 700)
    s_scope = get_sec("scope"); score_scope = _len_score(s_scope, 200, 700)
    s_resp = get_sec("responsibilit", "role")
    roles_found = set(m.group(0).lower() for m in RE_ROLES.finditer(s_resp))
    role_count_score = min(len(roles_found) / 4.0, 1.0)
    score_resp = 0.6 * _len_score(s_resp, 200, 700) + 0.4 * role_count_score
    s_proc = get_sec("procedure", "work instruction", "method", "steps")
    steps = RE_NUM_STEP.findall(s_proc); steps_score = min(len(steps) / 8.0, 1.0)
    score_proc = 0.6 * _len_score(s_proc, 600, 2000) + 0.4 * steps_score
    s_change = get_sec("change control", "revision history", "change management")
    has_terms = 1.0 if re.search(r'\bchange control\b|\brevision history\b|\bversion\b', s_change, re.I) else 0.0
    score_change = 0.5 * has_terms + 0.5 * _len_score(s_change, 150, 500)
    s_records = get_sec("record", "documentation", "documents", "forms", "logs")
    has_terms_n = len(re.findall(r'\brecord|log|form|signature|approval|retention|attachment|annexure\b', s_records, re.I))
    term_score = min(has_terms_n / 5.0, 1.0)
    id_bonus = 0.3 if RE_DOC_IDS.search(s_records) else 0.0
    score_records = min(0.7 * _len_score(s_records, 250, 800) + 0.3 * term_score + id_bonus, 1.0)
    s_alcoa = get_sec("data integrity", "alcoa")
    alcoa_found = sum(1 for pat in ALCOA_TERMS if re.search(pat, s_alcoa, re.I))
    score_alcoa = min(alcoa_found / 6.0, 1.0) * 0.5 + 0.5 * _len_score(s_alcoa, 120, 500)
    s_ref = get_sec("reference", "regulatory")
    ref_found = any(re.search(p, s_ref, re.I) for p in REFERENCE_TERMS)
    score_refs = max(1.0 if ref_found else 0.0, 0.5 * _len_score(s_ref, 80, 300))
    rows = {}; total_pct = 0.0
    for k, s in {"Purpose":score_purpose,"Scope":score_scope,"Responsibilities":score_resp,"Procedure":score_proc,"Change Control":score_change,"Records / Documentation":score_records,"Data Integrity (ALCOA)":score_alcoa,"Regulatory References":score_refs}.items():
        w = WEIGHTS[k]; pct = round(s*100, 1)
        rows[k] = {"Present/Depth (%)": pct, "Weight": w, "Weighted": round(pct*w, 1),
                   "Comments": "✓ Good" if pct >= 80 else ("△ Needs detail" if pct >= 40 else "✗ Weak/Missing")}
        total_pct += pct*w
    df = pd.DataFrame.from_dict(rows, orient="index"); return df, round(total_pct, 1)

# ---------------------------
# OpenAI helpers
# ---------------------------
_OPENAI_AVAILABLE_FLAG = _OPENAI_AVAILABLE

def _get_openai_key() -> Optional[str]:
    return ""
@st.cache_resource(show_spinner=False)
def _get_openai_client():
    if not _OPENAI_AVAILABLE_FLAG:
        return None, "OpenAI SDK not installed. Run: pip install openai"
    api_key = _get_openai_key()
    if not api_key:
        return None, "Missing OpenAI API key"
    # new SDK
    try:
        client = OpenAI(api_key=api_key)  # type: ignore
        return client, None
    except Exception:
        # legacy fallback
        try:
            openai.api_key = api_key  # type: ignore
            return "legacy", None
        except Exception as e:
            return None, f"Failed to init OpenAI: {e}"

def _chat_completion(messages: List[Dict], model: str = "gpt-4o-mini", temperature: float = 0.3, max_tokens: int = 1800) -> Tuple[Optional[str], Optional[str]]:
    client, err = _get_openai_client()
    if err: return None, err
    if client == "legacy":
        try:
            resp = openai.ChatCompletion.create(model=model, messages=messages, temperature=temperature, max_tokens=max_tokens)  # type: ignore
            return resp["choices"][0]["message"]["content"].strip(), None
        except Exception as e:
            return None, str(e)
    try:
        resp = client.chat.completions.create(model=model, messages=messages, temperature=temperature, max_tokens=max_tokens)  # type: ignore
        return resp.choices[0].message.content.strip(), None
    except Exception as e:
        return None, str(e)

def gpt_summarize_text(raw_text: str) -> Tuple[Optional[str], Optional[str]]:
    sys = {"role": "system", "content": "You are a pharma regulatory compliance assistant. Write concise, accurate summaries with regulatory context."}
    usr = {"role": "user", "content": f"Summarize this SOP. Focus on objective, scope, key steps, roles, records, and compliance touchpoints.\n\n{raw_text[:12000]}"}
    return _chat_completion([sys, usr], model="gpt-4o-mini", temperature=0.2, max_tokens=900)

def gpt_generate_full_sop(topic: str) -> Tuple[Optional[str], Optional[str]]:
    sys = {"role": "system", "content": "You write structured, cGMP-compliant SOPs for pharma manufacturing. Be formal and specific."}
    usr = {"role": "user", "content": ("Draft a detailed SOP in Markdown for the topic below. Target ~1,200–1,800 words (2–3 PDF pages). "
            "Use headings exactly in this order: Title, Purpose, Scope, Definitions, Responsibilities, Materials & Equipment, "
            "Safety & Precautions, Procedure (stepwise, numbered), Deviations & CAPA, Change Control, Records & Documentation, "
            "Data Integrity (ALCOA+), References, Revision History. Topic:\n" f"{topic.strip()}")}
    content, err = _chat_completion([sys, usr], model="gpt-4o-mini", temperature=0.3, max_tokens=3000)
    if err: return None, err
    return content, None

# ---------------------------
# Session state defaults
# ---------------------------
for key, default in [
    ("audit", []), ("summary", ""), ("summary_pdf", None), ("summary_docx", None),
    ("chat_history", []), ("generated_sop_md", ""), ("compliance_df", None),
    ("compliance_total", None), ("generated_sop_pdf", None), ("generated_sop_docx", None),
    ("jump_to_uploads", False),
]:
    if key not in st.session_state: st.session_state[key] = default

# ---------------------------
# Header & Sidebar
# ---------------------------
st.markdown("<h1 style='text-align:center; letter-spacing:1px;'>REGDOCGPT — Admin</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align:center; color:#5b6770; font-size:18px;'>Upload to Summarize • Chat to Generate SOPs • Audit Trail</p>", unsafe_allow_html=True)
st.markdown("---")

with st.sidebar:
    st.subheader("OpenAI status")
    key_present = True if _get_openai_key() else False
    st.write("API key:", "✅ found" if key_present else "❌ missing")
    st.markdown("---")
    st.write(f"**Logged in (admin):** {acct.get('username','admin')}")
    st.write(f"**Role:** {acct.get('role','admin').title()}")
    st.write(f"**Admin ID:** {acct.get('id','—')}")

# ---------------------------
# Layout
# ---------------------------
left, right = st.columns([1, 1])

# ---- LEFT: Upload & Summary (same as user page) ----
with left:
    if st.session_state.get("jump_to_uploads"):
        st.info("📂 My documents / uploads")
        st.session_state["jump_to_uploads"] = False

    st.markdown("### Upload SOP (auto-summarize)")
    uploaded = st.file_uploader("Browse files (.txt/.pdf/.docx)", type=["txt", "pdf", "docx"], label_visibility="collapsed")

    if uploaded is not None:
        content_bytes = uploaded.read()
        ext = os.path.splitext(uploaded.name)[1].lower()
        extracted_text = ""
        try:
            if ext == ".txt":
                extracted_text = content_bytes.decode("utf-8", errors="ignore")
            elif ext == ".pdf":
                if fitz is None:
                    st.warning("PyMuPDF not installed; cannot parse PDF.")
                else:
                    doc = fitz.open(stream=content_bytes, filetype="pdf")
                    extracted_text = "\n".join(page.get_text() for page in doc)
            elif ext == ".docx":
                if docx2txt is None:
                    st.warning("docx2txt not installed; cannot parse .docx.")
                else:
                    tmp = "_tmp_upload.docx"
                    with open(tmp, "wb") as f:
                        f.write(content_bytes)
                    try:
                        extracted_text = docx2txt.process(tmp)
                    finally:
                        try: os.remove(tmp)
                        except Exception: pass
        except Exception as e:
            st.error(f"Failed to extract text: {e}")

        with st.spinner("Summarizing with OpenAI…"):
            summary, err = gpt_summarize_text(extracted_text) if extracted_text.strip() else (None, "No text extracted")
            if err or not summary:
                st.info(f"Falling back to heuristic: {err or 'no content'}")
                summary = summarize_heuristic(extracted_text)

        st.session_state.summary = summary
        st.session_state.summary_pdf = make_pdf_from_text("SOP Summary", summary)
        st.session_state.summary_docx = make_docx_from_text("SOP Summary", summary)
        add_audit("Admin", "SOP uploaded", uploaded.name)
        add_audit("System", "SOP summarized", uploaded.name)

    st.markdown("### Summary")
    st.text_area("Current summary", value=st.session_state.summary or "", height=220)

    st.markdown("#### Download summary")
    disabled = not bool(st.session_state.summary)
    c1, c2 = st.columns(2)
    with c1:
        st.download_button("Download PDF", data=st.session_state.summary_pdf or b"", file_name="sop_summary.pdf",
                           mime="application/pdf", disabled=disabled, use_container_width=True,
                           on_click=lambda: add_audit("Admin", "Download", "sop_summary.pdf"))
    with c2:
        st.download_button("Download Word", data=st.session_state.summary_docx or b"", file_name="sop_summary.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           disabled=disabled, use_container_width=True,
                           on_click=lambda: add_audit("Admin", "Download", "sop_summary.docx"))

# ---- RIGHT: Chatbot & Generated SOP (same as user page) ----
with right:
    st.markdown("### Chatbot")
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])

    user_prompt = st.chat_input("Ask about the SOP, or say: generate sop for rapid mixer granulator…")
    if user_prompt:
        st.session_state.chat_history.append({"role": "user", "content": user_prompt})
        with st.chat_message("user"): st.write(user_prompt)

        q_lower = user_prompt.strip().lower()
        gen_match = re.search(r"(generate|create|draft)\s+sop\s+(for|on)\s+(.+)", q_lower)
        if gen_match:
            topic = user_prompt[gen_match.start(3):].strip()
            with st.spinner("Drafting SOP with OpenAI…"):
                sop_md, err = gpt_generate_full_sop(topic)
            if err or not sop_md:
                reply = f"Generation failed: {err or 'no content'}"
                st.session_state.chat_history.append({"role": "assistant", "content": reply})
                with st.chat_message("assistant"): st.write(reply)
            else:
                st.session_state.generated_sop_md = sop_md
                df, total = evaluate_sop_compliance(sop_md)
                st.session_state.compliance_df = df
                st.session_state.compliance_total = total
                st.session_state.generated_sop_pdf = make_pdf_from_text("Standard Operating Procedure", sop_md)
                st.session_state.generated_sop_docx = make_docx_from_text("Standard Operating Procedure", sop_md)
                add_audit("Assistant", "Generated SOP", topic[:120])
                with st.chat_message("assistant"):
                    st.write("Drafted SOP and calculated compliance score below. Scroll to **Generated SOP & Downloads**.")
        else:
            if not st.session_state.summary:
                reply = "I don’t see an uploaded SOP yet. Please upload a document to summarize first, or ask me to generate a new SOP (e.g., 'generate sop for rapid mixer granulator')."
            else:
                s = st.session_state.summary
                if any(k in q_lower for k in ["summary", "overview", "tl;dr", "what is this", "what does it say"]):
                    reply = f"Here’s the current summary of the SOP:\n\n{s}"
                elif any(k in q_lower for k in ["next step", "action", "what should", "checklist"]):
                    reply = ("Based on the SOP summary, suggested next steps:\n"
                             "1) Validate responsible roles & approvals\n"
                             "2) Confirm training requirements\n"
                             "3) Verify effective/retirement dates\n"
                             "4) Ensure change control links and related SOPs are referenced\n"
                             "5) Publish and notify impacted teams")
                elif "risk" in q_lower:
                    reply = ("From the summary, check for:\n"
                             "- Deviations and CAPA handling\n"
                             "- Data integrity controls (ALCOA+)\n"
                             "- Change control and impact assessment\n"
                             "- Batch record review steps\n"
                             "If any are absent, flag as a potential gap.")
                else:
                    reply = ("Here’s what I can infer from the current SOP summary:\n\n"
                             f"{s}\n\n"
                             "You can also ask me to *generate* a new SOP, e.g., 'generate sop for rapid mixer granulator'.")
            st.session_state.chat_history.append({"role": "assistant", "content": reply})
            add_audit("Assistant", "Replied", reply[:180] + ("…" if len(reply) > 180 else ""))
            with st.chat_message("assistant"): st.write(reply)

    if st.session_state.generated_sop_md:
        st.markdown("---"); st.markdown("### Generated SOP & Downloads")
        st.text_area("SOP (Markdown preview)", value=st.session_state.generated_sop_md, height=240)
        if st.session_state.compliance_df is not None:
            st.metric("Compliance Quality Score", f"{st.session_state.compliance_total}%")
            st.dataframe(st.session_state.compliance_df, use_container_width=True)
            st.download_button("Download Compliance Breakdown (CSV)",
                               data=st.session_state.compliance_df.to_csv().encode("utf-8"),
                               file_name="sop_compliance.csv", mime="text/csv")
        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("Download SOP (Markdown)",
                               data=st.session_state.generated_sop_md.encode("utf-8"),
                               file_name="generated_sop.md", use_container_width=True)
        with c2:
            st.download_button("Download SOP (PDF)",
                               data=st.session_state.generated_sop_pdf or b"",
                               file_name="generated_sop.pdf", mime="application/pdf",
                               use_container_width=True)
        with c3:
            st.download_button("Download SOP (Word)",
                               data=st.session_state.generated_sop_docx or b"",
                               file_name="generated_sop.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True)

# ---------------------------
# Audit Trail — Admin only
# ---------------------------
def _load_audits_df(actor: Optional[str], user_id: Optional[str], event: Optional[str],
                    search: Optional[str], limit: int = 2000) -> pd.DataFrame:
    """
    Pulls audits via sp_get_audits.
    Tries new signature: (@actor,@actor_role,@admin_id,@user_id,@event,@search,@limit,@offset)
    Falls back to old signature: (@actor,@user_id,@event,@search,@limit,@offset)
    Returns DataFrame with columns: ts, actor, actor_role, admin_id, user_id, event, detail
    """
    uid_int = None
    try:
        uid_int = int(user_id) if user_id else None
    except Exception:
        uid_int = None

    rows = []
    try:
        rows = db_get_audits(
            actor=actor or None,
            actor_role=None,
            admin_id=None,
            user_id=uid_int,
            event=event or None,
            search=search or None,
            limit=limit,
            offset=0,
        )
    except TypeError:
        try:
            rows = db_get_audits(
                actor=actor or None,
                user_id=uid_int,
                event=event or None,
                search=search or None,
                limit=limit,
                offset=0,
            )
        except Exception:
            rows = []
    except Exception:
        rows = []

    if not rows:
        mem = st.session_state.get("audit", [])
        if not mem:
            return pd.DataFrame(columns=["ts", "actor", "actor_role", "admin_id", "user_id", "event", "detail"])
        df = pd.DataFrame(mem).rename(columns={
            "Timestamp": "ts",
            "Actor": "actor",
            "UserID": "user_id",
            "Event": "event",
            "Detail": "detail",
        })
        for c in ["actor_role", "admin_id"]:
            if c not in df.columns:
                df[c] = None
        return df[["ts", "actor", "actor_role", "admin_id", "user_id", "event", "detail"]]

    df = pd.DataFrame(rows)
    for c in ["ts", "actor", "actor_role", "admin_id", "user_id", "event", "detail"]:
        if c not in df.columns:
            df[c] = None
    return df[["ts", "actor", "actor_role", "admin_id", "user_id", "event", "detail"]]

@st.cache_data(show_spinner=False, ttl=10)
def _fetch_all_for_filters(limit: int = 5000) -> pd.DataFrame:
    return _load_audits_df(actor=None, user_id=None, event=None, search=None, limit=limit)

st.markdown("---"); st.markdown("### Audit Trail")
if _bootstrap_err:
    st.caption(f"Note: Schema bootstrap warning earlier: {_bootstrap_err}")

base_df = _fetch_all_for_filters(limit=3000)
if base_df.empty and not st.session_state.get("audit"):
    st.caption("No audit events yet.")
else:
    col1, col2, col3, col4, col5 = st.columns([1, 1, 1, 1, 2])
    with col1:
        actor_opt = sorted([x for x in base_df["actor"].dropna().unique()])
        actor_sel = st.multiselect("Actor", options=actor_opt, default=actor_opt)
    with col2:
        role_opt = sorted([x for x in base_df["actor_role"].dropna().unique()])
        role_sel = st.multiselect("Actor Role", options=role_opt, default=role_opt)
    with col3:
        user_opt = sorted([x for x in base_df["user_id"].dropna().unique()])
        user_sel = st.multiselect("User ID", options=user_opt, default=user_opt)
    with col4:
        admin_opt = sorted([x for x in base_df["admin_id"].dropna().unique()])
        admin_sel = st.multiselect("Admin ID", options=admin_opt, default=admin_opt)
    with col5:
        event_opt = sorted([x for x in base_df["event"].dropna().unique()])
        event_sel = st.multiselect("Event", options=event_opt, default=event_opt)

    search = st.text_input("Search detail", value="", placeholder="Contains…")

    df = _load_audits_df(actor=None, user_id=None, event=None, search=search or None, limit=3000)

    if actor_sel:
        df = df[df["actor"].isin(actor_sel)]
    if role_sel:
        df = df[df["actor_role"].isin(role_sel)]
    if user_sel:
        df = df[df["user_id"].isin(user_sel)]
    if admin_sel:
        df = df[df["admin_id"].isin(admin_sel)]
    if event_sel:
        df = df[df["event"].isin(event_sel)]

    df_display = df.rename(columns={
        "ts": "Timestamp",
        "actor": "Actor",
        "actor_role": "Role",
        "admin_id": "AdminID",
        "user_id": "UserID",
        "event": "Event",
        "detail": "Detail",
    })

    st.dataframe(
        df_display[["Timestamp", "Actor", "Role", "AdminID", "UserID", "Event", "Detail"]],
        hide_index=True,
        use_container_width=True
    )

    csv = df_display.to_csv(index=False).encode("utf-8")
    st.download_button("Export audit CSV", data=csv, file_name="audit_trail.csv", mime="text/csv")
