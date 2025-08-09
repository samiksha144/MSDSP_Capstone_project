import io
import time
from datetime import datetime
import streamlit as st

# ---------- Page config ----------
st.set_page_config(page_title="REGDOCGPT", page_icon="📄", layout="wide")

# ---------- Utilities ----------
def now_iso():
    # Match the screenshot’s format: YYYY-MM-DD HH:MM:SS
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def add_audit(actor: str, event: str, detail: str = ""):
    """Append a row to the audit trail."""
    st.session_state.audit.insert(
        0,
        {"Timestamp": now_iso(), "Actor": actor, "Event": event, "Detail": detail},
    )

# Back-compat wrapper to keep your old calls working
def add_action(action: str):
    add_audit("System", action, "")

def summarize_sop(bytes_data) -> str:
    # Placeholder summary logic: take first 400 chars of text if decodable, else generic
    try:
        text = bytes_data.decode("utf-8", errors="ignore")
    except Exception:
        text = ""
    if not text:
        return "The SOP document outlines the procedures for regulatory compliance and quality management."
    cleaned = " ".join(text.split())
    return (cleaned[:350] + "…") if len(cleaned) > 350 else cleaned

def make_docx(summary: str) -> bytes:
    # Simple Word file using python-docx if available; otherwise fallback to .txt disguised as .docx
    try:
        from docx import Document  # type: ignore
        doc = Document()
        doc.add_heading("SOP Summary", level=1)
        doc.add_paragraph(summary)
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()
    except Exception:
        return f"SOP Summary\n\n{summary}".encode("utf-8")

def make_pdf(summary: str) -> bytes:
    # Simple PDF using reportlab if available; fallback to plain text bytes with .pdf name
    try:
        from reportlab.lib.pagesizes import LETTER  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
        bio = io.BytesIO()
        c = canvas.Canvas(bio, pagesize=LETTER)
        width, height = LETTER
        textobj = c.beginText(72, height - 72)
        textobj.setFont("Helvetica", 12)
        textobj.textLine("SOP Summary")
        textobj.textLine("")
        # Wrap summary lines
        import textwrap
        for line in textwrap.wrap(summary, width=88):
            textobj.textLine(line)
        c.drawText(textobj)
        c.showPage()
        c.save()
        return bio.getvalue()
    except Exception:
        return f"SOP Summary\n\n{summary}".encode("utf-8")

def bot_reply(user_msg: str) -> str:
    """Very simple deterministic assistant that answers based on the current summary."""
    s = (st.session_state.summary or "").strip()
    if not s:
        return "I don’t see an SOP uploaded yet. Please upload a document, then ask me about its summary or next steps."
    # Tiny heuristic “RAG”: if user asks for summary/overview, return the summary;
    # otherwise try to be helpful with the summary as context.
    q = user_msg.lower()
    if any(k in q for k in ["summary", "overview", "tl;dr", "what is this", "what does it say"]):
        return f"Here’s the current summary of the SOP:\n\n{s}"
    if any(k in q for k in ["next step", "action", "what should", "checklist"]):
        return (
            "Based on the SOP summary, suggested next steps:\n"
            "1) Validate responsible roles & approvals\n"
            "2) Confirm training requirements\n"
            "3) Verify effective/retirement dates\n"
            "4) Ensure change control links and related SOPs are referenced\n"
            "5) Publish and notify impacted teams"
        )
    if "risk" in q:
        return (
            "From the summary, check for:\n"
            "- Deviations and CAPA handling\n"
            "- Data integrity controls (ALCOA+)\n"
            "- Change control and impact assessment\n"
            "- Batch record review steps\n"
            "If any are absent, flag as a potential gap."
        )
    # Default: echo with context
    return (
        "Here’s what I can infer from the current SOP summary:\n\n"
        f"{s}\n\n"
        "If you upload the full text (or a richer extract), I can answer more specific questions."
    )

# ---------- Session state ----------
if "audit" not in st.session_state:
    # New, richer audit trail
    st.session_state.audit = []
if "actions" not in st.session_state:
    st.session_state.actions = []  # kept for compatibility (not displayed)
if "summary" not in st.session_state:
    st.session_state.summary = ""
if "uploaded_name" not in st.session_state:
    st.session_state.uploaded_name = None
if "pdf_bytes" not in st.session_state:
    st.session_state.pdf_bytes = None
if "docx_bytes" not in st.session_state:
    st.session_state.docx_bytes = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []  # [{"role":"user/assistant", "content": "..."}]

# ---------- Header ----------
st.markdown("<h1 style='text-align:center; letter-spacing:1px;'>REGDOCGPT</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align:center; color:#5b6770; font-size:18px;'>Generative AI Assistant for Regulatory Compliance in Pharma</p>", unsafe_allow_html=True)
st.markdown("---")

# ---------- Layout ----------
left, right = st.columns([1, 1])

with left:
    st.markdown("### Upload SOP")
    uploaded = st.file_uploader("Browse files", type=None, label_visibility="collapsed")

    if uploaded is not None:
        # Read content and cache in session
        content = uploaded.read()
        st.session_state.uploaded_name = uploaded.name
        add_audit("User", "SOP uploaded", uploaded.name)
        # Simulate processing
        with st.spinner("Summarizing SOP…"):
            time.sleep(0.6)
            st.session_state.summary = summarize_sop(content)
            st.session_state.pdf_bytes = make_pdf(st.session_state.summary)
            st.session_state.docx_bytes = make_docx(st.session_state.summary)
        add_audit("System", "SOP summarized", f"Source: {uploaded.name}")

with right:
    st.markdown("### Workspace")
    tab_summary, tab_chat = st.tabs(["Summary", "Chatbot"])

    with tab_summary:
        st.text_area(
            "The SOP document outlines the procedures for…",
            value=st.session_state.summary or "",
            height=180,
            label_visibility="collapsed",
        )

    with tab_chat:
        # Controls
        colA, colB = st.columns([1, 1])
        with colA:
            if st.button("Clear chat", use_container_width=True):
                st.session_state.chat_history = []
                add_audit("User", "Chat cleared")
        with colB:
            st.caption("Ask questions about the uploaded SOP.")

        # Render history
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        # Chat input (one per page; we keep it inside the tab for grouping)
        user_prompt = st.chat_input("Ask about the SOP…")
        if user_prompt:
            # Log + display user msg
            st.session_state.chat_history.append({"role": "user", "content": user_prompt})
            add_audit("User", "Asked", user_prompt)
            with st.chat_message("user"):
                st.write(user_prompt)

            # Bot reply
            reply = bot_reply(user_prompt)
            st.session_state.chat_history.append({"role": "assistant", "content": reply})
            add_audit("Assistant", "Replied", reply[:180] + ("…" if len(reply) > 180 else ""))
            with st.chat_message("assistant"):
                st.write(reply)

# ---------- Actions + Downloads ----------
st.markdown("---")
st.markdown("### Outputs")

c1, c2 = st.columns([3, 1])

with c1:
    st.caption("Generate files from the current summary.")
with c2:
    disabled = not bool(st.session_state.summary)

    def _on_pdf():
        add_audit("User", "Download", "sop_summary.pdf")

    def _on_docx():
        add_audit("User", "Download", "sop_summary.docx")

    st.download_button(
        label="Download PDF",
        data=st.session_state.pdf_bytes or b"",
        file_name="sop_summary.pdf",
        mime="application/pdf",
        disabled=disabled,
        use_container_width=True,
        on_click=_on_pdf,
    )
    st.download_button(
        label="Download Word",
        data=st.session_state.docx_bytes or b"",
        file_name="sop_summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        disabled=disabled,
        use_container_width=True,
        on_click=_on_docx,
    )

# ---------- Audit Trail ----------
st.markdown("---")
st.markdown("### Audit Trail")

import pandas as pd

audit_df = pd.DataFrame(st.session_state.audit)
if audit_df.empty:
    st.caption("No audit events yet.")
else:
    # Filters
    fcol1, fcol2, fcol3 = st.columns([1, 1, 2])
    with fcol1:
        actor_filter = st.multiselect(
            "Actor", options=sorted(audit_df["Actor"].unique()), default=list(sorted(audit_df["Actor"].unique()))
        )
    with fcol2:
        event_filter = st.multiselect(
            "Event", options=sorted(audit_df["Event"].unique()), default=list(sorted(audit_df["Event"].unique()))
        )
    with fcol3:
        search = st.text_input("Search detail", value="", placeholder="Contains…")

    filtered = audit_df[
        audit_df["Actor"].isin(actor_filter) & audit_df["Event"].isin(event_filter)
    ]
    if search.strip():
        s = search.strip().lower()
        filtered = filtered[filtered["Detail"].str.lower().str.contains(s, na=False)]

    st.dataframe(filtered, hide_index=True, use_container_width=True)

    # Export CSV of the *filtered* audit
    csv = filtered.to_csv(index=False).encode("utf-8")
    st.download_button(
        "Export audit CSV",
        data=csv,
        file_name="audit_trail.csv",
        mime="text/csv",
        use_container_width=False,
    )

# ---------- Footer hint ----------
st.markdown(
    "<div style='text-align:right; color:#98a2b3; font-size:12px;'></div>",
    unsafe_allow_html=True,
)
